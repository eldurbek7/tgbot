import os
import csv
import re
import asyncio
import logging
import zipfile
import time
from datetime import datetime, timezone, timedelta
from html import escape
from typing import Optional, Any, Awaitable, Callable, Dict

from asyncpg.exceptions import UniqueViolationError as PgIntegrityError
from cachetools import TTLCache

from aiogram import Bot, Dispatcher, F
from aiogram.enums import ChatMemberStatus
from aiogram.filters import Command
from aiogram.dispatcher.middlewares.base import BaseMiddleware
from aiogram.exceptions import TelegramBadRequest
from aiogram.types import (
    Message,
    CallbackQuery,
    InlineKeyboardMarkup,
    InlineKeyboardButton,
    FSInputFile,
)
from aiogram.utils.keyboard import InlineKeyboardBuilder

from db import (
    close_pool,
    execute_sync,
    fetch_sync,
    fetchrow_sync,
    init_pool,
    parse_rowcount,
    reset_request_sql_time,
    get_request_sql_time,
)

try:
    from openpyxl import Workbook
except ImportError:
    Workbook = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    Document = None

# =========================
# SOZLAMALAR
# =========================
BOT_TOKEN = os.getenv("BOT_TOKEN")
if not BOT_TOKEN:
    raise ValueError("BOT_TOKEN muhit o'zgaruvchisi topilmadi. Iltimos, BOT_TOKEN ni o'rnating.")

DATABASE_URL = os.getenv("DATABASE_URL")
if not DATABASE_URL:
    raise ValueError("DATABASE_URL muhit o'zgaruvchisi topilmadi. Iltimos, Railway PostgreSQL DATABASE_URL ni o'rnating.")

CHANNEL_USERNAME = os.getenv("CHANNEL_USERNAME", "@Qashqadaryo_PMM")
CHANNEL_URL = "https://t.me/Qashqadaryo_PMM"

ADMIN_IDS_RAW = os.getenv("ADMIN_IDS", "5298063089,7361393654")
ADMIN_IDS = [int(x.strip()) for x in ADMIN_IDS_RAW.split(",") if x.strip().isdigit()]

FACEBOOK_URL = "https://www.facebook.com/share/1E4ZVePTh4/"
INSTAGRAM_URL = "https://www.instagram.com/pedagogikmahorat"

DATA_DIR = os.getenv("DATA_DIR", "/app/data")
os.makedirs(DATA_DIR, exist_ok=True)

EXPORT_FILE = os.path.join(DATA_DIR, "votes_export.csv")
VOTES_XLSX_FILE = os.path.join(DATA_DIR, "votes_export.xlsx")
RATING_XLSX_FILE = os.path.join(DATA_DIR, "rating_export.xlsx")
USERS_XLSX_FILE = os.path.join(DATA_DIR, "users_export.xlsx")
COMPLAINTS_DOCX_FILE = os.path.join(DATA_DIR, "complaints_export.docx")
BACKUP_ZIP_FILE = os.path.join(DATA_DIR, "bot_backup.zip")
SUBJECTS_RANKING_XLSX_FILE = os.path.join(DATA_DIR, "subjects_ranking_export.xlsx")

SUBJECTS = {
    "s1": {
        "name": "Tillarni o'qitish metodikasi",
        "old_key": "tillarni_oqitish_metodikasi",
        "teachers": {
            "tom_1": "Norov Otajon Shomurodovich",
            "tom_2": "Abdixolikov Abdulazizxon Abduvohob o'g'li",
            "tom_3": "Azimova Nigora Anvar qizi",
            "tom_4": "Abatov Doston Ro'zimurod o'g'li",
            "tom_5": "Jalilova Komila Abdullayevna",
            "tom_6": "Oqboyeva Zulfiya Bobonazarovna",
            "tom_7": "Sevastyanova Nadejda Aleksandrovna",
            "tom_8": "Xidirova Feruza To'rayevna",
            "tom_9": "Ergasheva Dilorom Muradilloyevna",
        },
    },
    "s2": {
        "name": "Pedagogika, psixologiya va ta'lim menejmenti",
        "old_key": "pedagogika_psixologiya_va_talim_menejmenti",
        "teachers": {
            "pptm_1": "Umarov Lutfillo Murodilloyevich",
            "pptm_2": "Baratova Nasiba Turobovna",
            "pptm_3": "Bekmurodova Dilnoza Pirimovna",
            "pptm_4": "Meyliyev Lobar Nurmatovna",
            "pptm_5": "Ochilov Og'abek Narzullayevich",
            "pptm_6": "Shoniyozova Dilafruz Sabirovna",
            "pptm_7": "Yaratov Xamidjon Muxtorovich",
            "pptm_8": "Nazarov Asliddin Faxriddin o'g'li",
            "pptm_9": "Ergasheva Dilafruz Ergamqulovna",
            "pptm_10": "Soatov Asadulloh Jabborovich",
        },
    },
    "s3": {
        "name": "Aniq va tabiiy fanlar",
        "old_key": "aniq_va_tabiiy_fanlar",
        "teachers": {
            "atf_1": "Jobborov Farhod Bo'riyevich",
            "atf_2": "Karimova Habiba Abduraxmonovna",
            "atf_3": "Quldoshova Maftuna Jumanzar qizi",
            "atf_4": "Mallaev Xamro Ro'ziboyevich",
            "atf_5": "Mamatov Bekzod Farxotovich",
            "atf_6": "Pardaeva Muqaddas Zafar qizi",
            "atf_7": "Parmanov Jahongir Rayhonovich",
            "atf_8": "Rahmatullayev Erkin Shokirovich",
            "atf_9": "Suyarov Zoir Shojmardonovich",
            "atf_10": "Tursunova Maftuna Sulton qizi",
            "atf_11": "Umarov Ibrohimxon Norxuja o'g'li",
            "atf_12": "Chariev Rashid Ravshanovich",
            "atf_13": "Elmurodov Sherdil Ergashyevich",
            "atf_14": "Eshmonov Laziz Norxo'rja o'g'li",
            "atf_15": "Karaeva Dilfuzaxon Mamasharipovna",
            "atf_16": "Salomova Madina Sodiq qizi",
        },
    },
    "s4": {
        "name": "Amaliy va ijtimoiy fanlar",
        "old_key": "amaliy_va_ijtimoiy_fanlar",
        "teachers": {
            "aif_1": "Yo'ldashev Bekmirza Elmurodovich",
            "aif_2": "Jabboborov Laziz Hamza o'g'li",
            "aif_3": "Nurmatov Samandar Fayratovich",
            "aif_4": "Batoshov Inatillo Kungirovich",
            "aif_5": "Rajabov Ruslan Bozorovich",
            "aif_6": "Sanaev Azamat Alponovich",
            "aif_7": "Shamsiev Jahongir Qulmurod o'g'li",
            "aif_8": "Xudoyberdiev Axrorboy Nabi o'g'li",
            "aif_9": "Xasanova Gulnora Qorshanbiyevna",
            "aif_10": "Eshnazarova Maziya Allanazarovna",
        },
    },
    "s5": {
        "name": "Maktabgacha, boshlang'ich va maxsus ta'lim",
        "old_key": "maktabgacha_boshlangich_va_maxsus_talim",
        "teachers": {
            "mbmt_1": "Irisova Sayyora Rajabovna",
            "mbmt_2": "Azizova Dilnoz Yo'ldoshevna",
            "mbmt_3": "G'oyimov Umar Eshmurodovich",
            "mbmt_4": "Ziyotova Madina Mansur qizi",
            "mbmt_5": "Karimova Umida Sharopovna",
            "mbmt_6": "Qarshiyeva Guzal Alimardonovna",
            "mbmt_7": "Qurbanova Xusnora Xudoyberdi qizi",
            "mbmt_8": "Rajabova Xurshida Hakimovna",
            "mbmt_9": "Razzaqova Dilnoza Akramovna",
            "mbmt_10": "Sadinova Marjona Akmal qizi",
            "mbmt_11": "Shaxmurodova Dilxaxon Almardanovna",
            "mbmt_12": "Ergasheva Xusniya Mirzoxid qizi",
            "mbmt_13": "Zaripova Muslima Qurbonovna",
        },
    },
}
OLD_TO_NEW_SUBJECT = {v["old_key"]: k for k, v in SUBJECTS.items()}

logging.basicConfig(level=logging.INFO)

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

# =========================
# CACHE / PERFORMANCE
# =========================
SETTINGS_CACHE: TTLCache = TTLCache(maxsize=256, ttl=60)
USER_PREFS_CACHE: TTLCache = TTLCache(maxsize=20000, ttl=300)
SUBSCRIPTION_CACHE: TTLCache = TTLCache(maxsize=20000, ttl=600)
VOTED_CACHE: TTLCache = TTLCache(maxsize=20000, ttl=300)
SUBJECTS_CACHE: TTLCache = TTLCache(maxsize=4, ttl=60)
STATS_CACHE: TTLCache = TTLCache(maxsize=512, ttl=60)

def invalidate_stats_cache() -> None:
    STATS_CACHE.clear()

def invalidate_subjects_cache() -> None:
    SUBJECTS_CACHE.clear()
    invalidate_stats_cache()

class PerformanceMiddleware(BaseMiddleware):
    async def __call__(
        self,
        handler: Callable[[Any, Dict[str, Any]], Awaitable[Any]],
        event: Any,
        data: Dict[str, Any],
    ) -> Any:
        reset_request_sql_time()
        started = time.perf_counter()
        try:
            return await handler(event, data)
        finally:
            total_time = time.perf_counter() - started
            sql_time = get_request_sql_time()
            if isinstance(event, CallbackQuery):
                logging.info(
                    "BUTTON=%s SQL_TIME=%.4fs TOTAL_TIME=%.4fs",
                    event.data or "unknown",
                    sql_time,
                    total_time,
                )

dp.callback_query.middleware(PerformanceMiddleware())

db_lock = asyncio.Lock()
WAITING_COMPLAINT_TEXT = set()
COMPLAINT_STATE = {}
SUGGESTION_SPAM_LIMIT = 5
SUGGESTION_BREAK_SECONDS = 300
ADMIN_MANAGE_STATE = {}
COMPLAINT_MAX_LENGTH = 1000

LAST_REFRESH = {}
REFRESH_BUSY = set()
REFRESH_COOLDOWN_SECONDS = 1.5

# =========================
# O'ZBEKISTON VAQTI
# =========================
UZ_TZ = timezone(timedelta(hours=5), name="Asia/Tashkent")

def uz_now() -> datetime:
    return datetime.now(UZ_TZ)

# =========================
# POSTGRESQL ASYNCPG POOL WRAPPERS
# =========================
def _params(params):
    if params is None:
        return ()
    if isinstance(params, (list, tuple)):
        return tuple(params)
    return (params,)

def execute_query(sql: str, params=None, fetch: str = None):
    """Legacy-compatible helper backed by asyncpg global pool.

    fetch=None  -> execute and return affected row count
    fetch='one' -> return dict row
    fetch='all' -> return list[dict]
    """
    args = _params(params)
    if fetch == 'one':
        row = fetchrow_sync(sql, *args)
        return dict(row) if row else None
    if fetch == 'all':
        rows = fetch_sync(sql, *args)
        return [dict(r) for r in rows]
    status = execute_sync(sql, *args)
    return parse_rowcount(status)

def execute_query_plain(sql: str, params=None, fetch: str = None):
    """Tuple-style legacy helper backed by asyncpg global pool."""
    args = _params(params)
    if fetch == 'one':
        row = fetchrow_sync(sql, *args)
        return tuple(row) if row else None
    if fetch == 'all':
        rows = fetch_sync(sql, *args)
        return [tuple(r) for r in rows]
    status = execute_sync(sql, *args)
    return parse_rowcount(status)

# =========================
# LOTIN / KRILL
# =========================
def latin_to_cyrillic_text(text: str) -> str:
    pairs = [
        ("O\u2018", "\u040e"), ("o\u2018", "\u045e"),
        ("G\u2018", "\u0492"), ("g\u2018", "\u0493"),
        ("O'", "\u040e"), ("o'", "\u045e"),
        ("G'", "\u0492"), ("g'", "\u0493"),
        ("Sh", "\u0428"), ("sh", "\u0448"),
        ("Ch", "\u0427"), ("ch", "\u0447"),
        ("Ya", "\u042f"), ("ya", "\u044f"),
        ("Yo", "\u0401"), ("yo", "\u0451"),
        ("Yu", "\u042e"), ("yu", "\u044e"),
        ("Ts", "\u0426"), ("ts", "\u0446"),
    ]
    for old, new in pairs:
        text = text.replace(old, new)

    table = str.maketrans({
        "A": "\u0410", "a": "\u0430",
        "B": "\u0411", "b": "\u0431",
        "D": "\u0414", "d": "\u0434",
        "E": "\u0415", "e": "\u0435",
        "F": "\u0424", "f": "\u0444",
        "G": "\u0413", "g": "\u0433",
        "H": "\u04b2", "h": "\u04b3",
        "I": "\u0418", "i": "\u0438",
        "J": "\u0416", "j": "\u0436",
        "K": "\u041a", "k": "\u043a",
        "L": "\u041b", "l": "\u043b",
        "M": "\u041c", "m": "\u043c",
        "N": "\u041d", "n": "\u043d",
        "O": "\u041e", "o": "\u043e",
        "P": "\u041f", "p": "\u043f",
        "Q": "\u049a", "q": "\u049b",
        "R": "\u0420", "r": "\u0440",
        "S": "\u0421", "s": "\u0441",
        "T": "\u0422", "t": "\u0442",
        "U": "\u0423", "u": "\u0443",
        "V": "\u0412", "v": "\u0432",
        "X": "\u0425", "x": "\u0445",
        "Y": "\u0419", "y": "\u0439",
        "Z": "\u0417", "z": "\u0437",
        "`": "\u044a", "'": "\u044a", "\u2019": "\u044a",
    })
    return text.translate(table)


def translit_html_safe(text: str, script: str) -> str:
    parts = re.split(r"(<[^>]+>)", text)
    result = []
    for part in parts:
        if part.startswith("<") and part.endswith(">"):
            result.append(part)
        else:
            result.append(latin_to_cyrillic_text(part) if script == "cyrillic" else part)
    return "".join(result)


def get_user_script(user_id: int) -> str:
    cached = USER_PREFS_CACHE.get(user_id)
    if cached:
        script = cached.get("script", "latin")
        return script if script in ("latin", "cyrillic") else "latin"
    row = execute_query_plain("SELECT script, access_granted FROM user_prefs WHERE user_id = $1", (user_id,), fetch='one')
    if row:
        USER_PREFS_CACHE[user_id] = {"script": row[0] or "latin", "access_granted": int(row[1] or 0)}
        return row[0] if row[0] in ("latin", "cyrillic") else "latin"
    return "latin"


def set_user_script(user_id: int, script: str):
    if script not in ("latin", "cyrillic"):
        script = "latin"
    ensure_user(user_id)
    execute_query("UPDATE user_prefs SET script = $1 WHERE user_id = $2", (script, user_id))
    cached = USER_PREFS_CACHE.get(user_id, {"access_granted": 0})
    USER_PREFS_CACHE[user_id] = {"script": script, "access_granted": int(cached.get("access_granted", 0))}


def tr(user_id: int, text: str) -> str:
    return translit_html_safe(text, get_user_script(user_id))

def like_label(user_id: int) -> str:
    return "Лайк" if get_user_script(user_id) == "cyrillic" else "Like"

def dislike_label(user_id: int) -> str:
    return "Дислайк" if get_user_script(user_id) == "cyrillic" else "Dislike"


# =========================
# DB
# =========================
def normalize_subject_key(subject_key: str) -> str:
    if subject_key == "general":
        return "general"
    return OLD_TO_NEW_SUBJECT.get(subject_key, subject_key)

def apply_indexes():
    with open(os.path.join(os.path.dirname(__file__), "migrations_indexes.sql"), "r", encoding="utf-8") as f:
        execute_query(f.read())


def load_settings_cache() -> None:
    SETTINGS_CACHE.clear()
    rows = execute_query_plain("SELECT key, value FROM settings", fetch='all') or []
    for key, value in rows:
        SETTINGS_CACHE[key] = value

def init_db():
    execute_query("""
        CREATE TABLE IF NOT EXISTS votes (
            user_id BIGINT PRIMARY KEY,
            full_name TEXT,
            username TEXT,
            subject_key TEXT NOT NULL,
            teacher_key TEXT NOT NULL,
            voted_at TEXT
        );

        CREATE TABLE IF NOT EXISTS settings (
            key TEXT PRIMARY KEY,
            value TEXT
        );

        CREATE TABLE IF NOT EXISTS user_prefs (
            user_id BIGINT PRIMARY KEY,
            script TEXT DEFAULT 'latin',
            access_granted INTEGER DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS teacher_ratings (
            user_id BIGINT NOT NULL,
            full_name TEXT,
            username TEXT,
            subject_key TEXT NOT NULL,
            teacher_key TEXT NOT NULL,
            rating TEXT NOT NULL CHECK(rating IN ('like', 'dislike')),
            rated_at TEXT,
            PRIMARY KEY (user_id, subject_key, teacher_key)
        );

        CREATE TABLE IF NOT EXISTS complaints (
            id SERIAL PRIMARY KEY,
            user_id BIGINT NOT NULL,
            full_name TEXT,
            username TEXT,
            type TEXT DEFAULT 'general',
            subject_key TEXT,
            teacher_key TEXT,
            message_text TEXT NOT NULL,
            created_at TEXT
        );

        CREATE TABLE IF NOT EXISTS db_subjects (
            subject_key TEXT PRIMARY KEY,
            subject_name TEXT NOT NULL,
            sort_order INTEGER DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS db_teachers (
            teacher_key TEXT NOT NULL,
            subject_key TEXT NOT NULL,
            teacher_name TEXT NOT NULL,
            student_count INTEGER DEFAULT 0,
            PRIMARY KEY (subject_key, teacher_key)
        );
    """)

    # Mavjud jadvallar uchun user_id ustunini BIGINT ga o'tkazish
    for alter_sql in [
        "ALTER TABLE votes ALTER COLUMN user_id TYPE BIGINT",
        "ALTER TABLE user_prefs ALTER COLUMN user_id TYPE BIGINT",
        "ALTER TABLE teacher_ratings ALTER COLUMN user_id TYPE BIGINT",
        "ALTER TABLE complaints ALTER COLUMN user_id TYPE BIGINT",
    ]:
        try:
            execute_query(alter_sql)
        except Exception:
            pass

    migrate_old_subject_keys()
    _sync_subjects_to_db()
    apply_indexes()
    load_settings_cache()
    defaults = {
        "voting_open": "1",
        "auto_voting_enabled": "0",
        "auto_voting_start": "09:00",
        "auto_voting_end": "18:00",
    }
    for key, value in defaults.items():
        if get_setting(key, "") == "":
            set_setting(key, value)
    load_settings_cache()
    get_subjects_from_db(force=True)


def _sync_subjects_to_db():
    row = execute_query_plain("SELECT COUNT(*) FROM db_subjects", fetch='one')
    count = row[0] if row else 0
    if count == 0:
        for i, (skey, sdata) in enumerate(SUBJECTS.items()):
            execute_query(
                "INSERT INTO db_subjects (subject_key, subject_name, sort_order) VALUES ($1, $2, $3) ON CONFLICT DO NOTHING",
                (skey, sdata["name"], i)
            )
            for tkey, tname in sdata["teachers"].items():
                execute_query(
                    "INSERT INTO db_teachers (teacher_key, subject_key, teacher_name) VALUES ($1, $2, $3) ON CONFLICT DO NOTHING",
                    (tkey, skey, tname)
                )
    invalidate_subjects_cache()


def get_subjects_from_db(force: bool = False) -> dict:
    if not force and "subjects" in SUBJECTS_CACHE:
        return SUBJECTS_CACHE["subjects"]
    rows = execute_query_plain("SELECT subject_key, subject_name FROM db_subjects ORDER BY sort_order, subject_key", fetch='all')
    subjects = {}
    if rows:
        for skey, sname in rows:
            subjects[skey] = {"name": sname, "teachers": {}}
    teacher_rows = execute_query_plain(
        "SELECT subject_key, teacher_key, teacher_name, COALESCE(student_count, 0) FROM db_teachers ORDER BY subject_key, teacher_key",
        fetch='all'
    ) or []
    for skey, tkey, tname, student_count in teacher_rows:
        if skey not in subjects:
            subjects[skey] = {"name": skey, "teachers": {}}
        subjects[skey]["teachers"][tkey] = tname
        subjects[skey].setdefault("student_counts", {})[tkey] = int(student_count or 0)
    SUBJECTS_CACHE["subjects"] = subjects
    return subjects


def migrate_old_subject_keys():
    for old_key, new_key in OLD_TO_NEW_SUBJECT.items():
        execute_query("UPDATE votes SET subject_key = $1 WHERE subject_key = $2", (new_key, old_key))
        execute_query("UPDATE teacher_ratings SET subject_key = $1 WHERE subject_key = $2", (new_key, old_key))


def get_setting(key: str, default: str = "") -> str:
    if key in SETTINGS_CACHE:
        return SETTINGS_CACHE.get(key, default)
    row = execute_query_plain("SELECT value FROM settings WHERE key = $1", (key,), fetch='one')
    if row:
        SETTINGS_CACHE[key] = row[0]
        return row[0]
    return default


def set_setting(key: str, value: str):
    execute_query(
        "INSERT INTO settings (key, value) VALUES ($1, $2) ON CONFLICT (key) DO UPDATE SET value = EXCLUDED.value",
        (key, value)
    )
    SETTINGS_CACHE[key] = value


def ensure_user(user_id: int):
    if user_id in USER_PREFS_CACHE:
        return
    execute_query(
        "INSERT INTO user_prefs (user_id, script, access_granted) VALUES ($1, 'latin', 0) ON CONFLICT (user_id) DO NOTHING",
        (user_id,)
    )
    row = execute_query_plain("SELECT script, access_granted FROM user_prefs WHERE user_id = $1", (user_id,), fetch='one')
    if row:
        USER_PREFS_CACHE[user_id] = {"script": row[0] or "latin", "access_granted": int(row[1] or 0)}
    else:
        USER_PREFS_CACHE[user_id] = {"script": "latin", "access_granted": 0}


def has_access(user_id: int) -> bool:
    ensure_user(user_id)
    cached = USER_PREFS_CACHE.get(user_id, {})
    return bool(cached.get("access_granted", 0))


def require_access_only(user_id: int) -> bool:
    return has_access(user_id)


def grant_access(user_id: int):
    ensure_user(user_id)
    execute_query("UPDATE user_prefs SET access_granted = 1 WHERE user_id = $1", (user_id,))
    cached = USER_PREFS_CACHE.get(user_id, {"script": "latin"})
    USER_PREFS_CACHE[user_id] = {"script": cached.get("script", "latin"), "access_granted": 1}
    SUBSCRIPTION_CACHE[user_id] = True


def reset_access(user_id: int):
    ensure_user(user_id)
    execute_query("UPDATE user_prefs SET access_granted = 0 WHERE user_id = $1", (user_id,))
    cached = USER_PREFS_CACHE.get(user_id, {"script": "latin"})
    USER_PREFS_CACHE[user_id] = {"script": cached.get("script", "latin"), "access_granted": 0}
    SUBSCRIPTION_CACHE[user_id] = False


def is_admin(user_id: int) -> bool:
    return user_id in ADMIN_IDS


def is_voting_open() -> bool:
    return get_setting("voting_open", "1") == "1"


def open_voting():
    set_setting("voting_open", "1")


def close_voting():
    set_setting("voting_open", "0")


def is_auto_voting_enabled() -> bool:
    return get_setting("auto_voting_enabled", "0") == "1"


def set_auto_voting_enabled(enabled: bool):
    set_setting("auto_voting_enabled", "1" if enabled else "0")


def get_auto_voting_start() -> str:
    return get_setting("auto_voting_start", "09:00")


def get_auto_voting_end() -> str:
    return get_setting("auto_voting_end", "18:00")


def parse_hhmm(value: str) -> Optional[tuple[int, int]]:
    value = (value or "").strip()
    m = re.fullmatch(r"([01]?\d|2[0-3]):([0-5]\d)", value)
    if not m:
        return None
    return int(m.group(1)), int(m.group(2))


def set_auto_voting_schedule(start_time: str, end_time: str) -> bool:
    if not parse_hhmm(start_time) or not parse_hhmm(end_time):
        return False
    set_setting("auto_voting_start", start_time.strip())
    set_setting("auto_voting_end", end_time.strip())
    return True


def should_auto_voting_be_open(now: Optional[datetime] = None) -> bool:
    now = now or uz_now()
    start = parse_hhmm(get_auto_voting_start()) or (9, 0)
    end = parse_hhmm(get_auto_voting_end()) or (18, 0)
    current_minutes = now.hour * 60 + now.minute
    start_minutes = start[0] * 60 + start[1]
    end_minutes = end[0] * 60 + end[1]

    if start_minutes == end_minutes:
        return True  # 00:00-00:00 kabi bir xil vaqt 24 soat ochiq deb olinadi
    if start_minutes < end_minutes:
        return start_minutes <= current_minutes < end_minutes
    return current_minutes >= start_minutes or current_minutes < end_minutes


def apply_auto_voting_status() -> bool:
    """Avtomatik rejim yoqilgan bo'lsa, vaqtga qarab ovozni ochadi/yopadi."""
    if not is_auto_voting_enabled():
        return False
    target_open = should_auto_voting_be_open()
    current_open = is_voting_open()
    if target_open != current_open:
        set_setting("voting_open", "1" if target_open else "0")
        logging.info("Auto voting status changed: %s", "open" if target_open else "closed")
    return target_open


async def auto_voting_scheduler():
    while True:
        try:
            apply_auto_voting_status()
        except Exception as e:
            logging.exception("Auto voting scheduler xatosi: %s", e)
        await asyncio.sleep(30)


def has_voted(user_id: int) -> bool:
    cached = VOTED_CACHE.get(user_id)
    if cached is not None:
        return bool(cached)
    row = execute_query_plain("SELECT 1 FROM votes WHERE user_id = %s", (user_id,), fetch='one')
    voted = row is not None
    VOTED_CACHE[user_id] = voted
    return voted


def save_vote(user_id: int, full_name: str, username: str, subject_key: str, teacher_key: str) -> bool:
    subject_key = normalize_subject_key(subject_key)
    try:
        execute_query(
            """
            INSERT INTO votes (user_id, full_name, username, subject_key, teacher_key, voted_at)
            VALUES ($1, $2, $3, $4, $5, $6)
            """,
            (user_id, full_name, username, subject_key, teacher_key, uz_now().strftime("%Y-%m-%d %H:%M:%S"))
        )
        invalidate_stats_cache()
        VOTED_CACHE[user_id] = True
        return True
    except PgIntegrityError:
        VOTED_CACHE[user_id] = True
        return False


def get_total_votes(subject_key: Optional[str] = None) -> int:
    cache_key = f"total_votes:{normalize_subject_key(subject_key) if subject_key else 'all'}"
    if cache_key in STATS_CACHE:
        return int(STATS_CACHE[cache_key])
    if subject_key:
        row = execute_query_plain(
            "SELECT COUNT(*) FROM votes WHERE subject_key = $1",
            (normalize_subject_key(subject_key),), fetch='one'
        )
    else:
        row = execute_query_plain("SELECT COUNT(*) FROM votes", fetch='one')
    value = int(row[0] or 0) if row else 0
    STATS_CACHE[cache_key] = value
    return value



def get_vote_counts_map() -> dict[tuple[str, str], int]:
    cache_key = "vote_counts_map"
    if cache_key in STATS_CACHE:
        return STATS_CACHE[cache_key]
    rows = execute_query_plain(
        "SELECT subject_key, teacher_key, COUNT(*) FROM votes GROUP BY subject_key, teacher_key",
        fetch='all'
    ) or []
    result = {(normalize_subject_key(subject_key), teacher_key): int(count or 0) for subject_key, teacher_key, count in rows}
    STATS_CACHE[cache_key] = result
    return result

def reset_votes():
    execute_query("DELETE FROM votes")
    VOTED_CACHE.clear()
    invalidate_stats_cache()


def reset_ratings():
    execute_query("DELETE FROM teacher_ratings")
    invalidate_stats_cache()


def reset_complaints():
    execute_query("DELETE FROM complaints")
    invalidate_stats_cache()


def get_subject_name(subject_key: str) -> str:
    subject_key = normalize_subject_key(subject_key or "")
    subjects = get_subjects_from_db()
    if subject_key in subjects:
        return subjects[subject_key].get("name", subject_key)
    return SUBJECTS.get(subject_key, {}).get("name", subject_key)


def get_teacher_name(subject_key: str, teacher_key: str) -> str:
    subject_key = normalize_subject_key(subject_key or "")
    subjects = get_subjects_from_db()
    if subject_key in subjects:
        name = subjects[subject_key].get("teachers", {}).get(teacher_key)
        if name:
            return name
    return SUBJECTS.get(subject_key, {}).get("teachers", {}).get(teacher_key, teacher_key)


def build_progress_bar(percent: float, length: int = 14) -> str:
    filled = round((percent / 100) * length)
    filled = max(0, min(filled, length))
    return "▓" * filled + "░" * (length - filled)


def get_teacher_student_count(subject_key: str, teacher_key: str) -> int:
    subject_key = normalize_subject_key(subject_key or "")
    subjects = get_subjects_from_db()
    if subject_key in subjects:
        return int(subjects[subject_key].get("student_counts", {}).get(teacher_key, 0) or 0)
    return 0


def set_teacher_student_count(subject_key: str, teacher_key: str, student_count: int) -> bool:
    subject_key = normalize_subject_key(subject_key)
    student_count = max(0, int(student_count))
    rowcount = execute_query(
        "UPDATE db_teachers SET student_count = $1 WHERE subject_key = $2 AND teacher_key = $3",
        (student_count, subject_key, teacher_key)
    )
    invalidate_subjects_cache()
    return rowcount > 0


def get_teacher_participation_percent(subject_key: str, teacher_key: str, vote_count: int) -> float:
    student_count = get_teacher_student_count(subject_key, teacher_key)
    return get_vote_percent(vote_count, student_count) if student_count else 0.0


def get_all_teachers_flat():
    items = []
    subjects = get_subjects_from_db()
    for subject_key, subject_data in subjects.items():
        for teacher_key, teacher_name in subject_data["teachers"].items():
            items.append((subject_key, teacher_key, teacher_name))
    return items


def get_active_subjects() -> dict:
    return get_subjects_from_db()


# =========================
# RATING DB / STATS
# =========================
def save_teacher_rating(user_id: int, full_name: str, username: str, subject_key: str, teacher_key: str, rating: str):
    subject_key = normalize_subject_key(subject_key)
    execute_query(
        """
        INSERT INTO teacher_ratings (user_id, full_name, username, subject_key, teacher_key, rating, rated_at)
        VALUES (%s, %s, %s, %s, %s, %s, %s)
        ON CONFLICT (user_id, subject_key, teacher_key)
        DO UPDATE SET
            full_name = EXCLUDED.full_name,
            username = EXCLUDED.username,
            rating = EXCLUDED.rating,
            rated_at = EXCLUDED.rated_at
        """,
        (user_id, full_name, username, subject_key, teacher_key, rating, uz_now().strftime("%Y-%m-%d %H:%M:%S"))
    )
    invalidate_stats_cache()


def get_user_teacher_rating(user_id: int, subject_key: str, teacher_key: str) -> Optional[str]:
    row = execute_query_plain(
        "SELECT rating FROM teacher_ratings WHERE user_id = %s AND subject_key = %s AND teacher_key = %s",
        (user_id, normalize_subject_key(subject_key), teacher_key), fetch='one'
    )
    return row[0] if row else None


def get_rating_counts(subject_key: str, teacher_key: str):
    row = execute_query_plain(
        """
        SELECT
            SUM(CASE WHEN rating = 'like' THEN 1 ELSE 0 END),
            SUM(CASE WHEN rating = 'dislike' THEN 1 ELSE 0 END),
            COUNT(*)
        FROM teacher_ratings
        WHERE subject_key = %s AND teacher_key = %s
        """,
        (normalize_subject_key(subject_key), teacher_key), fetch='one'
    )
    like_count = int(row[0] or 0) if row else 0
    dislike_count = int(row[1] or 0) if row else 0
    total = int(row[2] or 0) if row else 0
    like_percent = (like_count / total * 100) if total else 0
    dislike_percent = (dislike_count / total * 100) if total else 0
    return like_count, dislike_count, total, like_percent, dislike_percent


def rating_rows():
    cache_key = "rating_rows"
    if cache_key in STATS_CACHE:
        return STATS_CACHE[cache_key]
    db_rows = execute_query_plain(
        """
        SELECT subject_key, teacher_key,
               SUM(CASE WHEN rating = 'like' THEN 1 ELSE 0 END) AS likes,
               SUM(CASE WHEN rating = 'dislike' THEN 1 ELSE 0 END) AS dislikes,
               COUNT(*) AS total
        FROM teacher_ratings
        GROUP BY subject_key, teacher_key
        """,
        fetch='all'
    ) or []
    counts = {(normalize_subject_key(s), t): (int(l or 0), int(d or 0), int(total or 0)) for s, t, l, d, total in db_rows}
    rows = []
    for subject_key, teacher_key, teacher_name in get_all_teachers_flat():
        like_count, dislike_count, total = counts.get((subject_key, teacher_key), (0, 0, 0))
        like_percent = (like_count / total * 100) if total else 0
        dislike_percent = (dislike_count / total * 100) if total else 0
        rows.append({
            "subject_key": subject_key,
            "subject_name": get_subject_name(subject_key),
            "teacher_key": teacher_key,
            "teacher_name": teacher_name,
            "like": like_count,
            "dislike": dislike_count,
            "total": total,
            "like_percent": like_percent,
            "dislike_percent": dislike_percent,
        })
    STATS_CACHE[cache_key] = rows
    return rows


def get_vote_percent(count: int, denominator: int) -> float:
    return (count / denominator * 100) if denominator > 0 else 0.0


def save_complaint(user_id: int, full_name: str, username: str, message_text: str, complaint_type: str = "general", subject_key: str = "", teacher_key: str = ""):
    execute_query(
        """
        INSERT INTO complaints (user_id, full_name, username, type, subject_key, teacher_key, message_text, created_at)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """,
        (
            user_id,
            full_name,
            username,
            complaint_type,
            normalize_subject_key(subject_key) if subject_key else "",
            teacher_key or "",
            message_text,
            uz_now().strftime("%Y-%m-%d %H:%M:%S")
        )
    )
    invalidate_stats_cache()


def get_last_complaint_for_user(user_id: int):
    row = execute_query_plain(
        "SELECT message_text, created_at FROM complaints WHERE user_id = %s ORDER BY id DESC LIMIT 1",
        (user_id,), fetch='one'
    )
    return row


def get_recent_suggestion_count(user_id: int) -> int:
    since = (uz_now() - timedelta(seconds=SUGGESTION_BREAK_SECONDS)).strftime("%Y-%m-%d %H:%M:%S")
    row = execute_query_plain(
        "SELECT COUNT(*) FROM complaints WHERE user_id = %s AND type = 'suggestion' AND created_at >= %s",
        (user_id, since), fetch='one'
    )
    return int(row[0] or 0) if row else 0


def complaint_allowed(user_id: int, message_text: str, complaint_type: str = "general"):
    text = (message_text or "").strip()
    if len(text) > COMPLAINT_MAX_LENGTH:
        return False, f"Xabar juda uzun. Iltimos, {COMPLAINT_MAX_LENGTH} ta belgidan oshirmang."

    last = get_last_complaint_for_user(user_id)
    if last:
        last_text = last[0]
        if (last_text or "").strip() == text:
            return False, "Siz aynan shu xabarni avval yuborgansiz. Iltimos, takroriy xabar yubormang."

    if complaint_type == "suggestion":
        recent_count = get_recent_suggestion_count(user_id)
        if recent_count >= SUGGESTION_SPAM_LIMIT:
            return False, "Takliflar uchun spam limiti: 5 ta xabardan keyin 5 daqiqa tanaffus qiling."

    return True, ""


def get_complaints_rows(limit: Optional[int] = None):
    sql = """
        SELECT id, user_id, full_name, username, type, subject_key, teacher_key, message_text, created_at
        FROM complaints
        ORDER BY id DESC
    """
    params = ()
    if limit:
        sql += " LIMIT %s"
        params = (limit,)
    return execute_query_plain(sql, params, fetch='all') or []


def get_complaints_count() -> int:
    row = execute_query_plain("SELECT COUNT(*) FROM complaints", fetch='one')
    return row[0] if row else 0


def get_complaints_text(user_id: int) -> str:
    rows = get_complaints_rows(limit=30)
    total = get_complaints_count()

    if not rows:
        return tr(user_id, "📩 <b>Shikoyat va takliflar</b>\n\nHali hech qanday xabar kelmagan.")

    lines = [f"📩 <b>Shikoyat va takliflar</b>\n\nJami: {total} ta\nOxirgi {len(rows)} ta xabar:\n"]
    for i, (cid, uid, full_name, username, complaint_type, subject_key, teacher_key, message_text, created_at) in enumerate(rows, start=1):
        safe_name = escape(full_name or "Noma'lum")
        safe_username = escape(username or "")
        safe_message = escape(message_text or "")
        type_label = "O'qituvchi ustidan shikoyat" if complaint_type == "teacher_complaint" else "Taklif" if complaint_type == "suggestion" else "Murojaat"
        line = f"{i}. <b>{escape(type_label)}</b> — <b>{safe_name}</b>"
        if safe_username:
            line += f" (@{safe_username})"
        line += f"\n   ID: <code>{uid}</code>"
        if subject_key:
            line += f"\n   Kafedra: {escape(get_subject_name(subject_key))}"
        if teacher_key:
            line += f"\n   O'qituvchi: {escape(get_teacher_name(subject_key, teacher_key))}"
        line += f"\n   Sana: {escape(created_at or '')}"
        line += f"\n   Xabar: {safe_message}"
        lines.append(line)

    text = "\n\n".join(lines)
    if len(text) > 4000:
        text = text[:4000] + "\n\n... qisqartirildi"
    return tr(user_id, text)


def export_complaints_to_docx() -> str:
    if Document is None:
        path = os.path.join(DATA_DIR, "complaints_export.txt")
        rows = get_complaints_rows()
        with open(path, "w", encoding="utf-8") as f:
            f.write("Shikoyat va takliflar\n")
            f.write(f"Jami: {len(rows)} ta\n\n")
            for i, (cid, uid, full_name, username, complaint_type, subject_key, teacher_key, message_text, created_at) in enumerate(rows, start=1):
                name = full_name or "Noma'lum"
                uname = f"@{username}" if username else ""
                type_label = "O'qituvchi ustidan shikoyat" if complaint_type == "teacher_complaint" else "Taklif" if complaint_type == "suggestion" else "Murojaat"
                teacher = get_teacher_name(subject_key or "", teacher_key or "") if teacher_key else ""
                f.write(f"{i}. {name} ({uname})\n")
                f.write(f"ID: {uid}\nTuri: {type_label}\nO'qituvchi: {teacher}\nSana: {created_at or ''}\nXabar: {message_text or ''}\n\n")
        return path

    rows = get_complaints_rows()
    doc = Document()

    title = doc.add_heading("Shikoyat va takliflar", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Jami xabarlar soni: ").bold = True
    p.add_run(str(len(rows)))

    p = doc.add_paragraph()
    p.add_run("Yaratilgan sana: ").bold = True
    p.add_run(uz_now().strftime("%Y-%m-%d %H:%M:%S"))

    if not rows:
        doc.add_paragraph("Hali hech qanday shikoyat yoki taklif kelmagan.")
    else:
        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        headers = ["№", "F.I.Sh", "Username", "Telegram ID", "Turi", "O'qituvchi", "Sana", "Xabar"]
        for idx, header in enumerate(headers):
            run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
            run.bold = True

        for i, (cid, uid, full_name, username, complaint_type, subject_key, teacher_key, message_text, created_at) in enumerate(rows, start=1):
            cells = table.add_row().cells
            type_label = "O'qituvchi ustidan shikoyat" if complaint_type == "teacher_complaint" else "Taklif" if complaint_type == "suggestion" else "Murojaat"
            cells[0].text = str(i)
            cells[1].text = full_name or "Noma'lum"
            cells[2].text = f"@{username}" if username else ""
            cells[3].text = str(uid)
            cells[4].text = type_label
            cells[5].text = get_teacher_name(subject_key or "", teacher_key or "") if teacher_key else ""
            cells[6].text = created_at or ""
            cells[7].text = message_text or ""

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)

    doc.save(COMPLAINTS_DOCX_FILE)
    return COMPLAINTS_DOCX_FILE


def get_subscription_required_alert(user_id: int) -> str:
    return tr(user_id, "Avval Telegram kanalga obuna bo'ling va ✅ Tekshirish tugmasini bosing.")

# =========================
# TEXTS
# =========================
def get_welcome_text(user_id: int) -> str:
    return tr(
        user_id,
        "🚀 <b>Botdan foydalanish uchun quyidagilarni bajaring:</b>\n\n"
        "1️⃣ 📸 Instagram sahifaga o'ting\n"
        "2️⃣ 📘 Facebook sahifaga o'ting\n"
        "3️⃣ 📢 Telegram kanalga obuna bo'ling\n\n"
        "Telegram obunasi bot tomonidan tekshiriladi. Instagram/Facebook tugmalari URL sifatida berilgan.\n\n"
        "👇 Telegram kanalga obuna bo'lgach, <b>✅ Tekshirish</b> tugmasini bosing."
    )

def get_home_text(user_id: int) -> str:
    return tr(user_id, "🏠 <b>Bosh menyu</b>\n\nKerakli bo'limni tanlang:")

def get_help_text(user_id: int) -> str:
    return tr(
        user_id,
        "ℹ️ <b>Yordam</b>\n\n"
        "• Telegram kanalga obuna bo'ling\n"
        "• So'ng ✅ Tekshirish tugmasini bosing\n"
        "• Ovoz berish uchun kafedra tanlanadi\n"
        "• Keyin o'qituvchi tanlanadi\n"
        "• Har bir foydalanuvchi asosiy ovozni faqat 1 marta beradi\n"
        "• Natijalarni istalgan payt ko'rishingiz mumkin"
    )

def get_already_voted_text(user_id: int) -> str:
    return tr(user_id, "✅ <b>Siz allaqachon ovoz berib bo'lgansiz</b>\n\nQayta ovoz berish mumkin emas.")

def get_closed_text(user_id: int) -> str:
    return tr(user_id, "🔒 <b>Ovoz berish hozircha yopilgan</b>\n\nAdmin tomonidan ovoz berish vaqtincha to'xtatilgan.")

def get_subject_select_text(user_id: int) -> str:
    return tr(user_id, "🗂 <b>Kafedrani tanlang</b>\n\nQuyidagi bo'limlardan birini tanlang:")

def get_teacher_select_text(user_id: int, subject_key: str) -> str:
    return tr(user_id, f"{get_subject_name(subject_key)}\n\n<b>O'qituvchini tanlang:</b>")

def get_rating_select_text(user_id: int) -> str:
    return tr(user_id, "⭐️ <b>O'qituvchilarni baholash</b>\n\nAvval kafedrani tanlang:")

def get_rating_teacher_text(user_id: int, subject_key: str) -> str:
    return tr(user_id, f"⭐️ <b>{get_subject_name(subject_key)}</b>\n\nBaholash uchun o'qituvchini tanlang:")

def get_rate_text(user_id: int, subject_key: str, teacher_key: str) -> str:
    current = get_user_teacher_rating(user_id, subject_key, teacher_key)
    if current == "like":
        current_text = f"Hozirgi bahoyingiz: 👍 {like_label(user_id)}"
    elif current == "dislike":
        current_text = f"Hozirgi bahoyingiz: 👎 {dislike_label(user_id)}"
    else:
        current_text = "Hozirgi bahoyingiz: hali baho berilmagan"
    like_count, dislike_count, total, like_percent, dislike_percent = get_rating_counts(subject_key, teacher_key)
    return tr(
        user_id,
        f"⭐️ <b>O'qituvchini baholash</b>\n\n"
        f"<b>Kafedra:</b> {get_subject_name(subject_key)}\n"
        f"<b>O'qituvchi:</b> {get_teacher_name(subject_key, teacher_key)}\n\n"
        f"{current_text}\n\n"
        f"👍 {like_label(user_id)}: {like_count} ta ({like_percent:.1f}%)\n"
        f"👎 {dislike_label(user_id)}: {dislike_count} ta ({dislike_percent:.1f}%)\n"
        f"Jami: {total} ta\n\n"
        f"Bahoni tanlang yoki o'zgartiring:"
    )

def get_complaint_intro_text(user_id: int) -> str:
    return tr(
        user_id,
        "📩 <b>Shikoyat va takliflar</b>\n\n"
        "Kerakli bo'limni tanlang:"
    )

def get_complaint_teacher_subject_text(user_id: int) -> str:
    return tr(user_id, "👤 <b>O'qituvchi ustidan shikoyat</b>\n\nAvval kafedrani tanlang:")

def get_complaint_teacher_select_text(user_id: int, subject_key: str) -> str:
    return tr(user_id, f"👤 <b>{get_subject_name(subject_key)}</b>\n\nShikoyat qilinadigan o'qituvchini tanlang:")

def get_complaint_write_text(user_id: int, mode: str, subject_key: str = "", teacher_key: str = "") -> str:
    if mode == "teacher_complaint":
        return tr(user_id, f"✍️ <b>Shikoyat matnini yozing</b>\n\n<b>O'qituvchi:</b> {get_teacher_name(subject_key, teacher_key)}\n\nMatnni bitta xabar qilib yuboring.")
    return tr(user_id, "✍️ <b>Taklif matnini yozing</b>\n\nTaklifingizni bitta xabar qilib yuboring.")

def get_complaint_confirm_text(user_id: int, state: dict) -> str:
    mode = state.get("mode", "suggestion")
    text = escape(state.get("text", ""))
    if mode == "teacher_complaint":
        subject_key = state.get("subject_key", "")
        teacher_key = state.get("teacher_key", "")
        title = "O'qituvchi ustidan shikoyat"
        info = f"<b>Kafedra:</b> {escape(get_subject_name(subject_key))}\n<b>O'qituvchi:</b> {escape(get_teacher_name(subject_key, teacher_key))}\n"
    else:
        title = "Taklif"
        info = ""
    return tr(user_id, f"❓ <b>Yuborishni tasdiqlaysizmi?</b>\n\n<b>Turi:</b> {title}\n{info}\n<b>Matn:</b>\n{text}")

def get_complaint_saved_text(user_id: int) -> str:
    return tr(user_id, "✅ <b>Xabaringiz qabul qilindi.</b>\n\nRahmat, murojaatingiz adminlarga yuborildi.")

def get_results_menu_text(user_id: int, is_admin_view: bool = False) -> str:
    title = "Admin natijalar bo'limi" if is_admin_view else "Natijalar bo'limi"
    return tr(user_id, f"📊 <b>{title}</b>\n\nKerakli bo'limni tanlang:")

def get_admin_panel_text(user_id: int) -> str:
    status_text = "🟢 Ochiq" if is_voting_open() else "🔴 Yopiq"
    return tr(user_id, f"🎛 <b>Admin panel</b>\n\nVoting holati: {status_text}\nJami ovozlar: {get_total_votes()}")

def get_general_results_text(user_id: int) -> str:
    rows_db = execute_query_plain(
        "SELECT subject_key, teacher_key, COUNT(*) FROM votes GROUP BY subject_key, teacher_key",
        fetch='all'
    ) or []
    counts = {}
    for subject_key, teacher_key, count in rows_db:
        counts[(normalize_subject_key(subject_key), teacher_key)] = count

    total_votes = sum(counts.values())
    lines = ["📊 <b>Umumiy natijalar</b>\n"]

    for subject_key, teacher_key, teacher_name in get_all_teachers_flat():
        count = counts.get((subject_key, teacher_key), 0)
        student_count = get_teacher_student_count(subject_key, teacher_key)
        percent = get_teacher_participation_percent(subject_key, teacher_key, count)
        denom_text = f"/{student_count}" if student_count else "/0"
        lines.append(
            f"<b>{teacher_name}</b> — {get_subject_name(subject_key)}\n"
            f"<code>{build_progress_bar(percent)}</code>  <b>{percent:.1f}%</b>  •  {count}{denom_text}\n"
        )

    lines.append(f"🗳 <b>Jami ovozlar:</b> {total_votes}")
    lines.append(f"{'🟢' if is_voting_open() else '🔴'} <b>Holat:</b> {'Ochiq' if is_voting_open() else 'Yopiq'}")

    text = "\n".join(lines)
    if len(text) > 4000:
        text = text[:4000] + "\n\n... qisqartirildi"
    return tr(user_id, text)


def get_subject_results_text(user_id: int, subject_key: str) -> str:
    subject_key = normalize_subject_key(subject_key)
    if subject_key not in get_subjects_from_db():
        return tr(user_id, "Noto'g'ri kafedra.")

    rows_db = execute_query_plain(
        "SELECT teacher_key, COUNT(*) FROM votes WHERE subject_key = %s GROUP BY teacher_key",
        (subject_key,), fetch='all'
    ) or []
    subject_counts = {teacher_key: count for teacher_key, count in rows_db}

    total_votes_row = execute_query_plain("SELECT COUNT(*) FROM votes", fetch='one')
    total_votes = total_votes_row[0] if total_votes_row else 0
    subject_total = sum(subject_counts.values())

    lines = [f"📊 <b>{get_subject_name(subject_key)} bo'yicha natijalar</b>\n"]

    for teacher_key, teacher_name in get_subjects_from_db().get(subject_key, {"teachers": {}}).get("teachers", {}).items():
        count = subject_counts.get(teacher_key, 0)
        student_count = get_teacher_student_count(subject_key, teacher_key)
        percent = get_teacher_participation_percent(subject_key, teacher_key, count)
        denom_text = f"/{student_count}" if student_count else "/0"
        lines.append(
            f"<b>{teacher_name}</b>\n"
            f"<code>{build_progress_bar(percent)}</code>  <b>{percent:.1f}%</b>  •  {count}{denom_text}\n"
        )

    lines.append(f"🗳 <b>Ushbu kafedra ovozlari:</b> {subject_total}")
    lines.append(f"🗳 <b>Jami ovozlar:</b> {total_votes}")
    lines.append(f"{'🟢' if is_voting_open() else '🔴'} <b>Holat:</b> {'Ochiq' if is_voting_open() else 'Yopiq'}")

    text = "\n".join(lines)
    if len(text) > 4000:
        text = text[:4000] + "\n\n... qisqartirildi"
    return tr(user_id, text)


def get_rating_stats_text(user_id: int, subject_key: Optional[str] = None) -> str:
    rows = rating_rows()
    if subject_key and subject_key != "general":
        subject_key = normalize_subject_key(subject_key)
        rows = [r for r in rows if r["subject_key"] == subject_key]
        title = f"⭐️ <b>{get_subject_name(subject_key)} — baholash foizlari</b>\n"
    else:
        title = "⭐️ <b>Umumiy baholash foizlari</b>\n"
    lines = [title]
    for r in rows:
        lines.append(
            f"<b>{r['teacher_name']}</b> — {r['subject_name']}\n"
            f"👍 {r['like']} ta ({r['like_percent']:.1f}%) | "
            f"👎 {r['dislike']} ta ({r['dislike_percent']:.1f}%) | "
            f"Jami: {r['total']}\n"
        )
    text = "\n".join(lines)
    return tr(user_id, text[:4000] + ("\n\n... qisqartirildi" if len(text) > 4000 else ""))


def get_top_ratings_text(user_id: int) -> str:
    rows = [r for r in rating_rows() if r["total"] > 0]

    def line_items(items, percent_key: str, icon: str):
        if not items:
            return "Ma'lumot yo'q"
        return "\n".join([
            f"{i}. {r['teacher_name']} — {r['subject_name']} | {icon} {r[percent_key]:.1f}% | Jami: {r['total']}"
            for i, r in enumerate(items, 1)
        ])

    high_like = sorted(rows, key=lambda r: (r["like_percent"], r["total"]), reverse=True)[:10]
    high_dislike = sorted(rows, key=lambda r: (r["dislike_percent"], r["total"]), reverse=True)[:10]

    text = (
        "🏆 <b>TOP reytinglar</b>\n\n"
        "🔝 <b>TOP 10 eng baland like nisbati</b>\n" + line_items(high_like, "like_percent", "👍") + "\n\n"
        "🔻 <b>TOP 10 eng baland dislike nisbati</b>\n" + line_items(high_dislike, "dislike_percent", "👎")
    )
    return tr(user_id, text[:4000] + ("\n\n... qisqartirildi" if len(text) > 4000 else ""))


def get_top_votes_text(user_id: int) -> str:
    items = []
    counts = get_vote_counts_map()
    for subject_key, teacher_key, teacher_name in get_all_teachers_flat():
        vote_count = counts.get((subject_key, teacher_key), 0)
        student_count = get_teacher_student_count(subject_key, teacher_key)
        participation = get_teacher_participation_percent(subject_key, teacher_key, vote_count)
        items.append((subject_key, teacher_key, teacher_name, vote_count, student_count, participation))

    items.sort(key=lambda x: (x[5], x[3]), reverse=True)
    top_rows = items[:10]
    total_votes = get_total_votes()

    if not top_rows or all(r[3] == 0 for r in top_rows):
        return tr(user_id, "🥇 <b>TOP 10 — O'qituvchi natijalari</b>\n\nHali hech kim ovoz bermagan.")

    medals = ["🥇", "🥈", "🥉", "4️⃣", "5️⃣", "6️⃣", "7️⃣", "8️⃣", "9️⃣", "🔟"]
    lines = [f"🥇 <b>TOP 10 — O'quvchilari ovoz bergan foizi bo'yicha</b>\n\n🗳 Jami ovozlar: {total_votes}\n"]

    for i, (subject_key, teacher_key, teacher_name, vote_count, student_count, percent) in enumerate(top_rows):
        subject_name = get_subject_name(subject_key)
        bar = build_progress_bar(percent)
        medal = medals[i] if i < len(medals) else f"{i+1}."
        denom_text = f"{student_count} o'quvchidan" if student_count else "o'quvchilar soni kiritilmagan"
        lines.append(
            f"{medal} <b>{teacher_name}</b>\n"
            f"   📂 {subject_name}\n"
            f"   <code>{bar}</code> <b>{percent:.1f}%</b> • {vote_count} ta ovoz ({denom_text})\n"
        )

    text = "\n".join(lines)
    if len(text) > 4000:
        text = text[:4000] + "\n\n... qisqartirildi"
    return tr(user_id, text)


def get_top_votes_by_subject_text(user_id: int) -> str:
    subjects = get_subjects_from_db()
    total_votes = get_total_votes()
    counts = get_vote_counts_map()
    medals = ["🥇", "🥈", "🥉", "4️⃣", "5️⃣", "6️⃣", "7️⃣", "8️⃣", "9️⃣", "🔟"]

    subject_rows = []
    for subject_key, subject_data in subjects.items():
        teacher_results = []
        subject_vote_count = 0
        teachers = subject_data.get("teachers", {})
        for teacher_key in teachers.keys():
            vote_count = counts.get((subject_key, teacher_key), 0)
            subject_vote_count += vote_count
            teacher_results.append(get_teacher_participation_percent(subject_key, teacher_key, vote_count))

        avg_percent = round(sum(teacher_results) / len(teacher_results), 2) if teacher_results else 0.0
        subject_rows.append((subject_key, subject_data["name"], avg_percent, subject_vote_count, len(teacher_results)))

    subject_rows.sort(key=lambda x: (x[2], x[3]), reverse=True)

    lines = [
        "🏫 <b>Kafedralar TOP 10 — o'qituvchilar natijalarining o'rtachasi</b>\n",
        f"🗳 <b>Jami ovozlar:</b> {total_votes}\n",
    ]

    for rank, (subject_key, sname, avg_percent, subject_vote_count, teacher_count) in enumerate(subject_rows[:10]):
        medal = medals[rank] if rank < len(medals) else f"{rank + 1}."
        bar = build_progress_bar(avg_percent)
        lines.append(
            f"{medal} <b>{sname}</b>\n"
            f"   <code>{bar}</code> <b>{avg_percent:.1f}%</b>\n"
            f"   👨‍🏫 {teacher_count} ta o'qituvchi o'rtachasi • 🗳 {subject_vote_count} ta ovoz\n"
        )

    lines.append(f"{'🟢' if is_voting_open() else '🔴'} <b>Holat:</b> {'Ochiq' if is_voting_open() else 'Yopiq'}")

    text = "\n".join(lines)
    if len(text) > 4000:
        text = text[:4000] + "\n\n... qisqartirildi"
    return tr(user_id, text)


def get_users_text(user_id: int) -> str:
    rows = execute_query_plain(
        "SELECT user_id, full_name, username, subject_key, teacher_key, voted_at FROM votes ORDER BY voted_at DESC",
        fetch='all'
    ) or []
    if not rows:
        return tr(user_id, "👥 Hali hech kim ovoz bermagan.")
    lines = [f"👥 <b>Kim kimga ovoz berdi</b>\n\nJami: {len(rows)} ta foydalanuvchi\n"]
    for i, (uid, full_name, username, subject_key, teacher_key, voted_at) in enumerate(rows, start=1):
        subject_key = normalize_subject_key(subject_key)
        name = full_name or "Noma'lum"
        line = f"{i}. <b>{name}</b>"
        if username:
            line += f" (@{username})"
        line += f"\n   → Kafedra: {get_subject_name(subject_key)}"
        line += f"\n   → O'qituvchi: {get_teacher_name(subject_key, teacher_key)}"
        line += f"\n   → ID: <code>{uid}</code>"
        if voted_at:
            line += f"\n   → {voted_at}"
        lines.append(line)
    text = "\n\n".join(lines)
    return tr(user_id, text[:4000] + ("\n\n... qisqartirildi" if len(text) > 4000 else ""))


def get_my_vote_text(user_id: int) -> str:
    row = execute_query_plain(
        "SELECT subject_key, teacher_key, voted_at FROM votes WHERE user_id = %s",
        (user_id,), fetch='one'
    )
    if not row:
        return tr(user_id, "🧾 <b>Mening ovozim</b>\n\nSiz hali asosiy ovoz bermagansiz.")

    subject_key, teacher_key, voted_at = row[0], row[1], row[2]
    subject_key = normalize_subject_key(subject_key)
    sana = voted_at if voted_at else "Noma'lum"
    return tr(
        user_id,
        f"🧾 <b>Mening ovozim</b>\n\n"
        f"<b>Kafedra:</b> {get_subject_name(subject_key)}\n"
        f"<b>O'qituvchi:</b> {get_teacher_name(subject_key, teacher_key)}\n"
        f"<b>Sana:</b> {sana}"
    )


def get_my_ratings_text(user_id: int) -> str:
    rows = execute_query_plain(
        "SELECT subject_key, teacher_key, rating, rated_at FROM teacher_ratings WHERE user_id = %s ORDER BY rated_at DESC",
        (user_id,), fetch='all'
    ) or []
    if not rows:
        return tr(user_id, "⭐️ <b>Mening baholarim</b>\n\nSiz hali o'qituvchilarga like/dislike bermagansiz.")

    lines = [f"⭐️ <b>Mening baholarim</b>\n\nJami: {len(rows)} ta baho\n"]
    for i, (subject_key, teacher_key, rating, rated_at) in enumerate(rows, start=1):
        subject_key = normalize_subject_key(subject_key)
        icon = "👍" if rating == "like" else "👎"
        label = like_label(user_id) if rating == "like" else dislike_label(user_id)
        lines.append(
            f"{i}. <b>{get_teacher_name(subject_key, teacher_key)}</b>\n"
            f"   Kafedra: {get_subject_name(subject_key)}\n"
            f"   Baho: {icon} {label}\n"
            f"   Sana: {rated_at or ''}"
        )

    text = "\n\n".join(lines)
    return tr(user_id, text[:4000] + ("\n\n... qisqartirildi" if len(text) > 4000 else ""))


def get_teacher_detailed_stats_text(user_id: int, subject_key: str, teacher_key: str) -> str:
    subject_key = normalize_subject_key(subject_key)
    if subject_key not in get_subjects_from_db() or teacher_key not in get_subjects_from_db().get(subject_key, {}).get("teachers", {}):
        return tr(user_id, "Noto'g'ri o'qituvchi tanlandi.")

    vote_count = get_vote_counts_map().get((subject_key, teacher_key), 0)
    subject_total = get_total_votes(subject_key)
    total_votes = get_total_votes()
    vote_percent_subject = get_vote_percent(vote_count, subject_total)
    vote_percent_total = get_vote_percent(vote_count, total_votes)
    student_count = get_teacher_student_count(subject_key, teacher_key)
    participation_percent = get_teacher_participation_percent(subject_key, teacher_key, vote_count)

    return tr(
        user_id,
        f"👤 <b>O'qituvchi statistikasi</b>\n\n"
        f"<b>O'qituvchi:</b> {get_teacher_name(subject_key, teacher_key)}\n"
        f"<b>Kafedra:</b> {get_subject_name(subject_key)}\n\n"
        f"👥 <b>O'quvchilari soni:</b> {student_count} ta\n"
        f"🗳 <b>Ovoz berganlar:</b> {vote_count} ta\n"
        f"📈 <b>Ishtirok foizi:</b> {participation_percent:.1f}%\n"
        f"📊 <b>Kafedra ichidagi ulushi:</b> {vote_percent_subject:.1f}%\n"
        f"🌐 <b>Jami ovozlar ichidagi ulushi:</b> {vote_percent_total:.1f}%"
    )


def export_subjects_ranking_to_excel() -> str:
    total_votes = get_total_votes()
    subjects = get_subjects_from_db()
    counts = get_vote_counts_map()

    subject_rows = []
    for subject_key, subject_data in subjects.items():
        teacher_rows = []
        subject_vote_count = 0
        for teacher_key, teacher_name in subject_data.get("teachers", {}).items():
            vote_count = counts.get((subject_key, teacher_key), 0)
            student_count = get_teacher_student_count(subject_key, teacher_key)
            percent = round(get_teacher_participation_percent(subject_key, teacher_key, vote_count), 2)
            subject_vote_count += vote_count
            teacher_rows.append((teacher_key, teacher_name, vote_count, student_count, percent))
        avg_percent = round(sum(r[4] for r in teacher_rows) / len(teacher_rows), 2) if teacher_rows else 0.0
        subject_rows.append((subject_key, subject_data["name"], avg_percent, subject_vote_count, len(teacher_rows), teacher_rows))

    subject_rows.sort(key=lambda x: (x[2], x[3]), reverse=True)

    if Workbook is None:
        path = os.path.join(DATA_DIR, "subjects_ranking_export.csv")
        with open(path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(["O'rin", "Kafedra", "O'rtacha natija (%)", "Ovozlar soni", "O'qituvchilar soni"])
            for rank, (_, sname, avg_percent, vote_count, teacher_count, _teacher_rows) in enumerate(subject_rows, start=1):
                writer.writerow([rank, sname, avg_percent, vote_count, teacher_count])
        return path

    wb = Workbook()
    wb.remove(wb.active)

    ws = wb.create_sheet("Kafedralar reytingi")
    ws_append_header(ws, ["O'rin", "Kafedra", "O'rtacha natija (%)", "Ovozlar soni", "O'qituvchilar soni", "Jami ovozlar"])
    for rank, (subject_key, sname, avg_percent, vote_count, teacher_count, teacher_rows_data) in enumerate(subject_rows, start=1):
        ws.append([rank, sname, avg_percent, vote_count, teacher_count, total_votes if rank == 1 else ""])

    for subject_key, sname, avg_percent, subject_vote_count, teacher_count, teacher_rows_data in subject_rows:
        ws_sub = wb.create_sheet((sname[:28] + " natija")[:31])
        ws_append_header(ws_sub, ["O'rin", "O'qituvchi", "Ovozlar", "O'quvchilar soni", "Natija (%)"])
        teacher_rows_data.sort(key=lambda x: (x[4], x[2]), reverse=True)
        for rank2, (_teacher_key, tname, vote_count, student_count, percent) in enumerate(teacher_rows_data, start=1):
            ws_sub.append([rank2, tname, vote_count, student_count, percent])

    wb.save(SUBJECTS_RANKING_XLSX_FILE)
    return SUBJECTS_RANKING_XLSX_FILE


def create_backup_zip() -> str:
    votes_path = export_votes_to_excel()
    users_path = export_users_to_excel()
    complaints_path = export_complaints_to_docx()
    subjects_ranking_path = export_subjects_ranking_to_excel()

    with zipfile.ZipFile(BACKUP_ZIP_FILE, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in (votes_path, users_path, complaints_path, subjects_ranking_path):
            if path and os.path.exists(path):
                zf.write(path, arcname=os.path.basename(path))
    return BACKUP_ZIP_FILE


def get_results_text_by_scope(user_id: int, scope: str) -> str:
    scope = normalize_subject_key(scope)
    if scope == "general":
        return get_general_results_text(user_id)
    if scope in get_subjects_from_db():
        return get_subject_results_text(user_id, scope)
    return tr(user_id, "Noto'g'ri bo'lim.")

# =========================
# EXPORT
# =========================
def export_votes_to_csv() -> str:
    rows = execute_query_plain(
        "SELECT user_id, full_name, username, subject_key, teacher_key, voted_at FROM votes ORDER BY voted_at DESC",
        fetch='all'
    ) or []
    with open(EXPORT_FILE, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["User ID", "Full Name", "Username", "Subject", "Teacher", "Voted At"])
        for user_id, full_name, username, subject_key, teacher_key, voted_at in rows:
            subject_key = normalize_subject_key(subject_key)
            writer.writerow([user_id, full_name or "", username or "", get_subject_name(subject_key), get_teacher_name(subject_key, teacher_key), voted_at or ""])
    return EXPORT_FILE


def ws_append_header(ws, headers):
    ws.append(headers)
    for cell in ws[1]:
        cell.style = "Headline 4"


def export_votes_to_excel() -> str:
    if Workbook is None:
        return export_votes_to_csv()
    wb = Workbook()
    wb.remove(wb.active)
    counts = get_vote_counts_map()

    ws = wb.create_sheet("Umumiy ovozlar")
    ws_append_header(ws, ["User ID", "Full Name", "Username", "Kafedra", "O'qituvchi", "Voted At"])
    rows = execute_query_plain(
        "SELECT user_id, full_name, username, subject_key, teacher_key, voted_at FROM votes ORDER BY voted_at DESC",
        fetch='all'
    ) or []
    for user_id, full_name, username, subject_key, teacher_key, voted_at in rows:
        subject_key = normalize_subject_key(subject_key)
        ws.append([user_id, full_name or "", username or "", get_subject_name(subject_key), get_teacher_name(subject_key, teacher_key), voted_at or ""])

    for subject_key, subject_data in get_subjects_from_db().items():
        ws = wb.create_sheet(subject_data["name"][:31])
        ws_append_header(ws, ["User ID", "Full Name", "Username", "O'qituvchi", "Voted At"])
        srows = execute_query_plain(
            "SELECT user_id, full_name, username, teacher_key, voted_at FROM votes WHERE subject_key = %s ORDER BY voted_at DESC",
            (subject_key,), fetch='all'
        ) or []
        for user_id, full_name, username, teacher_key, voted_at in srows:
            ws.append([user_id, full_name or "", username or "", get_teacher_name(subject_key, teacher_key), voted_at or ""])

    ws = wb.create_sheet("Umumiy natija")
    ws_append_header(ws, ["O'rin (Reyting)", "Kafedra", "O'qituvchi", "Ovozlar", "O'quvchilar soni", "Ishtirok foizi"])
    all_teachers_data = []
    for subject_key, teacher_key, teacher_name in get_all_teachers_flat():
        count = counts.get((subject_key, teacher_key), 0)
        scount = get_teacher_student_count(subject_key, teacher_key)
        percent = round(get_teacher_participation_percent(subject_key, teacher_key, count), 2)
        all_teachers_data.append((get_subject_name(subject_key), teacher_name, count, scount, percent))
    all_teachers_data.sort(key=lambda x: (x[4], x[2]), reverse=True)
    for rank, (subject_name, teacher_name, count, scount, percent) in enumerate(all_teachers_data, start=1):
        ws.append([rank, subject_name, teacher_name, count, scount, percent])

    for subject_key, subject_data in get_subjects_from_db().items():
        ws = wb.create_sheet((subject_data["name"][:24] + " natija")[:31])
        ws_append_header(ws, ["O'rin (Reyting)", "O'qituvchi", "Ovozlar", "O'quvchilar soni", "Ishtirok foizi"])
        subject_teachers_data = []
        for teacher_key, teacher_name in subject_data["teachers"].items():
            count = counts.get((subject_key, teacher_key), 0)
            scount = get_teacher_student_count(subject_key, teacher_key)
            percent = round(get_teacher_participation_percent(subject_key, teacher_key, count), 2)
            subject_teachers_data.append((teacher_name, count, scount, percent))
        subject_teachers_data.sort(key=lambda x: (x[3], x[1]), reverse=True)
        for rank, (teacher_name, count, scount, percent) in enumerate(subject_teachers_data, start=1):
            ws.append([rank, teacher_name, count, scount, percent])

    wb.save(VOTES_XLSX_FILE)
    return VOTES_XLSX_FILE


def export_rating_to_excel() -> str:
    if Workbook is None:
        path = os.path.join(DATA_DIR, "rating_export.csv")
        with open(path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(["Kafedra", "O'qituvchi", "Like", "Dislike", "Jami", "Like %", "Dislike %"])
            for r in rating_rows():
                writer.writerow([r["subject_name"], r["teacher_name"], r["like"], r["dislike"], r["total"], round(r["like_percent"], 2), round(r["dislike_percent"], 2)])
        return path

    wb = Workbook()
    wb.remove(wb.active)

    ws = wb.create_sheet("Umumiy rating")
    ws_append_header(ws, ["O'rin (Reyting)", "Kafedra", "O'qituvchi", "Like", "Dislike", "Jami", "Like %", "Dislike %"])
    sorted_all = sorted(rating_rows(), key=lambda x: x["total"], reverse=True)
    for rank, r in enumerate(sorted_all, start=1):
        ws.append([rank, r["subject_name"], r["teacher_name"], r["like"], r["dislike"], r["total"], round(r["like_percent"], 2), round(r["dislike_percent"], 2)])

    for subject_key, subject_data in get_subjects_from_db().items():
        ws = wb.create_sheet((subject_data["name"][:24] + " rating")[:31])
        ws_append_header(ws, ["O'rin (Reyting)", "O'qituvchi", "Like", "Dislike", "Jami", "Like %", "Dislike %"])
        subject_rows = sorted([x for x in rating_rows() if x["subject_key"] == subject_key], key=lambda x: x["total"], reverse=True)
        for rank, r in enumerate(subject_rows, start=1):
            ws.append([rank, r["teacher_name"], r["like"], r["dislike"], r["total"], round(r["like_percent"], 2), round(r["dislike_percent"], 2)])

    ws = wb.create_sheet("Umumiy ovozlar")
    ws_append_header(ws, ["User ID", "Full Name", "Username", "Kafedra", "O'qituvchi", "Rating", "Rated At"])
    rrows = execute_query_plain(
        "SELECT user_id, full_name, username, subject_key, teacher_key, rating, rated_at FROM teacher_ratings ORDER BY rated_at DESC",
        fetch='all'
    ) or []
    for user_id, full_name, username, subject_key, teacher_key, rating, rated_at in rrows:
        subject_key = normalize_subject_key(subject_key)
        ws.append([user_id, full_name or "", username or "", get_subject_name(subject_key), get_teacher_name(subject_key, teacher_key), rating, rated_at or ""])

    wb.save(RATING_XLSX_FILE)
    return RATING_XLSX_FILE


def export_users_to_excel() -> str:
    if Workbook is None:
        path = os.path.join(DATA_DIR, "users_export.csv")
        rows = execute_query_plain(
            """
            SELECT up.user_id, up.script, up.access_granted,
                   v.full_name, v.username, v.subject_key, v.teacher_key, v.voted_at
            FROM user_prefs up
            LEFT JOIN votes v ON up.user_id = v.user_id
            ORDER BY up.user_id
            """,
            fetch='all'
        ) or []
        with open(path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(["User ID", "Til", "Kirish", "Ism", "Username", "Kafedra", "O'qituvchi", "Ovoz sanasi"])
            for row in rows:
                uid, script, access, full_name, username, subject_key, teacher_key, voted_at = row
                subject_key = normalize_subject_key(subject_key) if subject_key else ""
                writer.writerow([
                    uid, script, "Ha" if access else "Yo'q",
                    full_name or "", username or "",
                    get_subject_name(subject_key) if subject_key else "",
                    get_teacher_name(subject_key, teacher_key) if subject_key and teacher_key else "",
                    voted_at or ""
                ])
        return path

    wb = Workbook()
    wb.remove(wb.active)

    ws_all = wb.create_sheet("Barcha foydalanuvchilar")
    ws_append_header(ws_all, [
        "User ID", "Til", "Kirish", "Ism", "Username",
        "Kafedra", "O'qituvchi", "Ovoz sanasi"
    ])
    all_rows = execute_query_plain(
        """
        SELECT up.user_id, up.script, up.access_granted,
               v.full_name, v.username, v.subject_key, v.teacher_key, v.voted_at
        FROM user_prefs up
        LEFT JOIN votes v ON up.user_id = v.user_id
        ORDER BY up.user_id
        """,
        fetch='all'
    ) or []
    for uid, script, access, full_name, username, subject_key, teacher_key, voted_at in all_rows:
        subject_key = normalize_subject_key(subject_key) if subject_key else ""
        ws_all.append([
            uid,
            script or "latin",
            "Ha" if access else "Yo'q",
            full_name or "",
            username or "",
            get_subject_name(subject_key) if subject_key else "",
            get_teacher_name(subject_key, teacher_key) if subject_key and teacher_key else "",
            voted_at or ""
        ])

    ws_voted = wb.create_sheet("Ovoz berganlar")
    ws_append_header(ws_voted, [
        "User ID", "Ism", "Username", "Kafedra", "O'qituvchi", "Ovoz sanasi"
    ])
    voted_rows = execute_query_plain(
        "SELECT user_id, full_name, username, subject_key, teacher_key, voted_at FROM votes ORDER BY voted_at DESC",
        fetch='all'
    ) or []
    for uid, full_name, username, subject_key, teacher_key, voted_at in voted_rows:
        subject_key = normalize_subject_key(subject_key)
        ws_voted.append([
            uid,
            full_name or "",
            username or "",
            get_subject_name(subject_key),
            get_teacher_name(subject_key, teacher_key),
            voted_at or ""
        ])

    ws_not_voted = wb.create_sheet("Ovoz bermaganlar")
    ws_append_header(ws_not_voted, ["User ID", "Til", "Kirish"])
    not_voted_rows = execute_query_plain(
        """
        SELECT up.user_id, up.script, up.access_granted
        FROM user_prefs up
        LEFT JOIN votes v ON up.user_id = v.user_id
        WHERE v.user_id IS NULL
        ORDER BY up.user_id
        """,
        fetch='all'
    ) or []
    for uid, script, access in not_voted_rows:
        ws_not_voted.append([uid, script or "latin", "Ha" if access else "Yo'q"])

    wb.save(USERS_XLSX_FILE)
    return USERS_XLSX_FILE


def export_complaints_to_word() -> str:
    rows = execute_query_plain(
        "SELECT user_id, full_name, username, message_text, created_at FROM complaints ORDER BY created_at DESC",
        fetch='all'
    ) or []

    if Document is None:
        txt_path = os.path.join(DATA_DIR, "shikoyat_takliflar.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("Shikoyat va takliflar\n")
            f.write("=" * 30 + "\n\n")
            if not rows:
                f.write("Hali shikoyat yoki taklif yo'q.\n")
            for i, (user_id, full_name, username, message_text, created_at) in enumerate(rows, 1):
                name = full_name or "Noma'lum"
                uname = f"@{username}" if username else "yo'q"
                f.write(f"{i}. Foydalanuvchi: {name}\n")
                f.write(f"   Username: {uname}\n")
                f.write(f"   ID: {user_id}\n")
                f.write(f"   Sana: {created_at or ''}\n")
                f.write(f"   Matn: {message_text or ''}\n")
                f.write("-" * 30 + "\n")
        return txt_path

    doc = Document()
    doc.add_heading("Shikoyat va takliflar", level=1)

    if not rows:
        doc.add_paragraph("Hali shikoyat yoki taklif yo'q.")
    else:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["№", "F.I.Sh", "Username", "User ID", "Sana", "Matn"]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header

        for i, (user_id, full_name, username, message_text, created_at) in enumerate(rows, 1):
            cells = table.add_row().cells
            cells[0].text = str(i)
            cells[1].text = full_name or "Noma'lum"
            cells[2].text = f"@{username}" if username else ""
            cells[3].text = str(user_id)
            cells[4].text = created_at or ""
            cells[5].text = message_text or ""

    doc.save(COMPLAINTS_DOCX_FILE)
    return COMPLAINTS_DOCX_FILE

# =========================
# SUBSCRIPTION
# =========================
async def check_user_subscription(user_id: int, *, force: bool = False, allow_cached_on_error: bool = True) -> bool:
    if not force:
        cached = SUBSCRIPTION_CACHE.get(user_id)
        if cached is not None:
            return bool(cached)
        if has_access(user_id):
            SUBSCRIPTION_CACHE[user_id] = True
            return True

    try:
        member = await bot.get_chat_member(CHANNEL_USERNAME, user_id)
        if member.status in {ChatMemberStatus.CREATOR, ChatMemberStatus.ADMINISTRATOR, ChatMemberStatus.MEMBER}:
            SUBSCRIPTION_CACHE[user_id] = True
            return True
        if member.status == ChatMemberStatus.RESTRICTED:
            ok = bool(getattr(member, "is_member", False))
            SUBSCRIPTION_CACHE[user_id] = ok
            return ok
        SUBSCRIPTION_CACHE[user_id] = False
        return False
    except Exception as e:
        logging.error(f"Obunani tekshirishda xatolik: {e}")
        if allow_cached_on_error and has_access(user_id):
            SUBSCRIPTION_CACHE[user_id] = True
            return True
        return False


async def refresh_access_from_channel(user_id: int) -> bool:
    """Force-check channel membership and update local access cache.

    Used on voting-critical paths so a user who leaves the channel cannot vote
    using an old cached access flag.
    """
    ok = await check_user_subscription(user_id, force=True, allow_cached_on_error=False)
    if ok:
        await asyncio.to_thread(grant_access, user_id)
    else:
        await asyncio.to_thread(reset_access, user_id)
    return ok


async def safe_edit_message(callback: CallbackQuery, text: str, reply_markup: Optional[InlineKeyboardMarkup] = None):
    try:
        await callback.message.edit_text(text=text, parse_mode="HTML", reply_markup=reply_markup)
    except TelegramBadRequest as e:
        if "message is not modified" in str(e).lower():
            return
        logging.error(f"edit_text xatosi: {e}")
        try:
            await callback.message.answer(text=text, parse_mode="HTML", reply_markup=reply_markup)
        except Exception as e2:
            logging.error(f"answer fallback xatosi: {e2}")
    except Exception as e:
        logging.error(f"safe_edit_message umumiy xato: {e}")


def add_refresh_time(text: str, user_id: int) -> str:
    return text + tr(user_id, f"  ⏱ Yangilandi: {uz_now().strftime('%H:%M:%S.%f')[:-3]}")


def can_start_refresh(user_id: int, key: str) -> bool:
    now = uz_now().timestamp()
    refresh_key = (user_id, key)

    if refresh_key in REFRESH_BUSY:
        return False

    last = LAST_REFRESH.get(refresh_key, 0)
    if now - last < REFRESH_COOLDOWN_SECONDS:
        return False

    REFRESH_BUSY.add(refresh_key)
    LAST_REFRESH[refresh_key] = now
    return True


def finish_refresh(user_id: int, key: str):
    REFRESH_BUSY.discard((user_id, key))


def get_settings_text(user_id: int) -> str:
    script = get_user_script(user_id)
    current = "Lotin" if script == "latin" else "Крилл"
    return tr(
        user_id,
        f"⚙️ <b>Sozlamalar</b>\n\n"
        f"Hozirgi yozuv: <b>{current}</b>\n\n"
        f"Kerakli yozuv turini tanlang:"
    )

# =========================
# KEYBOARDS
# =========================
def subscription_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text="📸 Instagram", url=INSTAGRAM_URL))
    kb.row(InlineKeyboardButton(text="📘 Facebook", url=FACEBOOK_URL))
    kb.row(InlineKeyboardButton(text="📢 Telegram", url=CHANNEL_URL))
    kb.row(InlineKeyboardButton(text="✅ Tekshirish", callback_data="check_subscription"))
    return kb.as_markup()


def home_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    if has_access(user_id):
        kb.row(InlineKeyboardButton(text=tr(user_id, "🗳 Ovoz berish"), callback_data="go_vote_panel"))
        kb.row(InlineKeyboardButton(text=tr(user_id, "📊 Natijalar"), callback_data="show_results_menu_user"))
        kb.row(InlineKeyboardButton(text=tr(user_id, "🥇 TOP 10 natijalar"), callback_data="user_top_votes"))
        kb.row(InlineKeyboardButton(text=tr(user_id, "🧾 Mening ovozim"), callback_data="my_vote"))
        kb.row(InlineKeyboardButton(text=tr(user_id, "📩 Shikoyat va takliflar"), callback_data="go_complaint_panel"))
    else:
        kb.row(InlineKeyboardButton(text=tr(user_id, "✅ Obunani tekshirish"), callback_data="check_subscription"))

    kb.row(
        InlineKeyboardButton(text=tr(user_id, "ℹ️ Yordam"), callback_data="help_info"),
        InlineKeyboardButton(text=tr(user_id, "⚙️ Sozlamalar"), callback_data="user_settings")
    )
    return kb.as_markup()


def settings_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    current = get_user_script(user_id)
    latin_text = "✅ Lotin" if current == "latin" else "Lotin"
    cyrillic_text = "✅ Крилл" if current == "cyrillic" else "Крилл"
    kb.row(
        InlineKeyboardButton(text=latin_text, callback_data="set_script:latin"),
        InlineKeyboardButton(text=cyrillic_text, callback_data="set_script:cyrillic")
    )
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Orqaga"), callback_data="go_home"))
    return kb.as_markup()


def subjects_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    for subject_key, subject_data in get_subjects_from_db().items():
        kb.row(InlineKeyboardButton(text=tr(user_id, subject_data["name"]), callback_data=f"subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def teachers_keyboard(user_id: int, subject_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    teachers = list(get_subjects_from_db().get(subject_key, {"teachers": {}}).get("teachers", {}).items())
    for i in range(0, len(teachers), 2):
        row = [InlineKeyboardButton(text=tr(user_id, teacher_name), callback_data=f"vote:{subject_key}:{teacher_key}") for teacher_key, teacher_name in teachers[i:i + 2]]
        kb.row(*row)
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Kafedralarga qaytish"), callback_data="go_vote_panel"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def rating_subjects_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    for subject_key, subject_data in get_subjects_from_db().items():
        kb.row(InlineKeyboardButton(text=tr(user_id, subject_data["name"]), callback_data=f"rating_subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def rating_teachers_keyboard(user_id: int, subject_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    teachers = list(get_subjects_from_db().get(subject_key, {"teachers": {}}).get("teachers", {}).items())
    for i in range(0, len(teachers), 2):
        row = [InlineKeyboardButton(text=tr(user_id, teacher_name), callback_data=f"rating_teacher:{subject_key}:{teacher_key}") for teacher_key, teacher_name in teachers[i:i + 2]]
        kb.row(*row)
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Kafedralarga qaytish"), callback_data="go_rating_panel"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def rate_keyboard(user_id: int, subject_key: str, teacher_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=f"👍 {like_label(user_id)}", callback_data=f"rate:like:{subject_key}:{teacher_key}"),
        InlineKeyboardButton(text=f"👎 {dislike_label(user_id)}", callback_data=f"rate:dislike:{subject_key}:{teacher_key}")
    )
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ O'qituvchilar"), callback_data=f"rating_subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def results_menu_keyboard_user(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "🥇 TOP 10 natijalar"), callback_data="user_top_votes"))
    for subject_key, subject_data in get_subjects_from_db().items():
        kb.row(InlineKeyboardButton(text=tr(user_id, subject_data["name"]), callback_data=f"show_results_user:{subject_key}"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Orqaga"), callback_data="go_home"))
    return kb.as_markup()


def results_menu_keyboard_admin(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    for subject_key, subject_data in get_subjects_from_db().items():
        kb.row(InlineKeyboardButton(text=subject_data["name"], callback_data=f"show_results_admin:{subject_key}"))
    kb.row(InlineKeyboardButton(text="⬅️ Sozlamalar", callback_data="admin_settings_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def rating_results_menu_keyboard_admin(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    for subject_key, subject_data in get_subjects_from_db().items():
        kb.row(InlineKeyboardButton(text=subject_data["name"], callback_data=f"show_rating_stats:{subject_key}"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def results_keyboard_user(user_id: int, scope: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data=f"refresh_results_user:{scope}"),
        InlineKeyboardButton(text=tr(user_id, "📂 Bo'limlar"), callback_data="show_results_menu_user")
    )
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def results_keyboard_admin(user_id: int, scope: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data=f"refresh_results_admin:{scope}"),
        InlineKeyboardButton(text=tr(user_id, "📂 Bo'limlar"), callback_data="admin_results")
    )
    kb.row(InlineKeyboardButton(text="⬅️ Ovoz natijalari", callback_data="admin_results"))
    kb.row(InlineKeyboardButton(text="⬅️ Sozlamalar", callback_data="admin_settings_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def rating_stats_keyboard_admin(user_id: int, scope: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data=f"refresh_rating_stats:{scope}"),
        InlineKeyboardButton(text=tr(user_id, "📂 Bo'limlar"), callback_data="admin_rating_stats")
    )
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def after_vote_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "📊 Natijalar"), callback_data="show_results_menu_user"),
        InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home")
    )
    return kb.as_markup()


def confirm_vote_keyboard(user_id: int, subject_key: str, teacher_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "✅ Tasdiqlash"), callback_data=f"confirm_vote:{subject_key}:{teacher_key}"),
        InlineKeyboardButton(text=tr(user_id, "❌ Bekor qilish"), callback_data=f"subject:{subject_key}")
    )
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ O'qituvchilar"), callback_data=f"subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def simple_back_home_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Orqaga"), callback_data="go_home"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def teacher_stats_subjects_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    for subject_key, subject_data in get_subjects_from_db().items():
        kb.row(InlineKeyboardButton(text=tr(user_id, subject_data["name"]), callback_data=f"teacher_stats_subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def teacher_stats_teachers_keyboard(user_id: int, subject_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    teachers = list(get_subjects_from_db().get(subject_key, {"teachers": {}}).get("teachers", {}).items())
    for i in range(0, len(teachers), 2):
        row = [InlineKeyboardButton(text=tr(user_id, teacher_name), callback_data=f"teacher_stats:{subject_key}:{teacher_key}") for teacher_key, teacher_name in teachers[i:i + 2]]
        kb.row(*row)
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Kafedralarga qaytish"), callback_data="admin_teacher_stats"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def teacher_stats_keyboard(user_id: int, subject_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ O'qituvchilar"), callback_data=f"teacher_stats_subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def admin_panel_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text="🥇 TOP 10 natijalar", callback_data="admin_top_votes_menu"))
    kb.row(InlineKeyboardButton(text="👤 O'qituvchi statistikasi", callback_data="admin_teacher_stats"))
    kb.row(InlineKeyboardButton(text="📩 Shikoyat va takliflar", callback_data="admin_complaints_menu"))
    kb.row(InlineKeyboardButton(text="🧹 Tozalash", callback_data="admin_cleanup_menu"))
    kb.row(InlineKeyboardButton(text="📥 Ma'lumotlarni yuklab olish", callback_data="admin_download_menu"))
    kb.row(InlineKeyboardButton(text="⚙️ Sozlamalar", callback_data="admin_settings_menu"))
    return kb.as_markup()


def admin_cleanup_menu_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text="♻️ Ovozlarni tozalash", callback_data="admin_reset_votes_confirm"))
    kb.row(InlineKeyboardButton(text="🧹 Shikoyat/takliflarni tozalash", callback_data="admin_reset_complaints_confirm"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def admin_download_menu_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text="📁 Ovozlar Excel", callback_data="admin_export_votes_excel"))
    kb.row(InlineKeyboardButton(text="🏫 Kafedralar natijasi Excel", callback_data="admin_export_subjects_excel"))
    kb.row(InlineKeyboardButton(text="👥 Foydalanuvchilar Excel", callback_data="admin_export_users_excel"))
    kb.row(InlineKeyboardButton(text="👤 Shikoyatlar Excel", callback_data="admin_export_teacher_complaints_excel"))
    kb.row(InlineKeyboardButton(text="💡 Takliflar Excel", callback_data="admin_export_suggestions_excel"))
    kb.row(InlineKeyboardButton(text="💾 Backup", callback_data="admin_backup"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def admin_settings_menu_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text="🔓 Manual ochish", callback_data="admin_open"),
        InlineKeyboardButton(text="🔒 Manual yopish", callback_data="admin_close")
    )
    auto_button = "🔴 Avto rejimni o'chirish" if is_auto_voting_enabled() else "🟢 Avto rejimni yoqish"
    kb.row(InlineKeyboardButton(text=auto_button, callback_data="admin_auto_voting_toggle"))
    kb.row(InlineKeyboardButton(text="⏰ Avto vaqtni belgilash", callback_data="admin_auto_voting_set_time"))
    kb.row(InlineKeyboardButton(text="👥 Foydalanuvchilar", callback_data="admin_users"))
    kb.row(InlineKeyboardButton(text="⚙️ Bo'lim va o'qituvchilarni boshqarish", callback_data="admin_manage_menu"))
    kb.row(InlineKeyboardButton(text="📊 Ovoz natijalari", callback_data="admin_results"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def reset_confirm_keyboard(user_id: int, mode: str = "votes") -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    if mode == "votes":
        yes_cb = "admin_reset_votes"
    elif mode == "rating":
        yes_cb = "admin_reset_rating"
    elif mode == "complaints":
        yes_cb = "admin_reset_complaints"
    else:
        yes_cb = "admin_reset_votes"
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "❌ Bekor qilish"), callback_data="cancel_reset"),
        InlineKeyboardButton(text="✅ Ha, o'chirish", callback_data=yes_cb)
    )
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def complaint_menu_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "👤 O'qituvchi ustidan shikoyat yuborish"), callback_data="complaint_teacher_start"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "💡 Taklif yuborish"), callback_data="suggestion_start"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Orqaga"), callback_data="go_home"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def complaint_subjects_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    for subject_key, subject_data in get_subjects_from_db().items():
        kb.row(InlineKeyboardButton(text=tr(user_id, subject_data["name"]), callback_data=f"complaint_subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Orqaga"), callback_data="go_complaint_panel"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def complaint_teachers_keyboard(user_id: int, subject_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    teachers = list(get_subjects_from_db().get(subject_key, {"teachers": {}}).get("teachers", {}).items())
    for i in range(0, len(teachers), 2):
        row = [InlineKeyboardButton(text=tr(user_id, teacher_name), callback_data=f"complaint_teacher:{subject_key}:{teacher_key}") for teacher_key, teacher_name in teachers[i:i + 2]]
        kb.row(*row)
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Kafedralarga qaytish"), callback_data="complaint_teacher_start"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def complaint_cancel_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Orqaga"), callback_data="go_complaint_panel"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "❌ Bekor qilish"), callback_data="cancel_complaint"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def complaint_confirm_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "✅ Yuborish"), callback_data="confirm_complaint_send"),
        InlineKeyboardButton(text=tr(user_id, "❌ Bekor qilish"), callback_data="cancel_complaint")
    )
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Orqaga"), callback_data="go_complaint_panel"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def admin_complaints_menu_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text="👤 O'qituvchi ustidan shikoyatlar", callback_data="admin_complaints_teacher"))
    kb.row(InlineKeyboardButton(text="💡 Takliflar", callback_data="admin_complaints_suggestions"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def complaints_keyboard_admin(user_id: int, complaint_type: str = "all") -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    refresh_cb = f"refresh_admin_complaints:{complaint_type}"
    kb.row(InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data=refresh_cb))
    kb.row(InlineKeyboardButton(text="⬅️ Shikoyat va takliflar", callback_data="admin_complaints_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def users_keyboard_admin(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data="refresh_admin_users"),
        InlineKeyboardButton(text="📊 Excel", callback_data="admin_export_users_excel")
    )
    kb.row(InlineKeyboardButton(text="⬅️ Sozlamalar", callback_data="admin_settings_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def get_complaints_text_filtered(user_id: int, complaint_type: str) -> str:
    complaint_type = "suggestion" if complaint_type == "suggestion" else "teacher_complaint"
    title = "💡 <b>Takliflar</b>" if complaint_type == "suggestion" else "👤 <b>O'qituvchi ustidan shikoyatlar</b>"
    rows = execute_query_plain(
        """
        SELECT id, user_id, full_name, username, type, subject_key, teacher_key, message_text, created_at
        FROM complaints
        WHERE type = %s
        ORDER BY id DESC
        LIMIT 30
        """,
        (complaint_type,), fetch='all'
    ) or []
    total_row = execute_query_plain("SELECT COUNT(*) FROM complaints WHERE type = %s", (complaint_type,), fetch='one')
    total = total_row[0] if total_row else 0
    if not rows:
        return tr(user_id, f"{title}\n\nHali ma'lumot yo'q.")

    lines = [f"{title}\n\nJami: {total} ta\nOxirgi {len(rows)} ta xabar:\n"]
    for i, (cid, uid, full_name, username, _ctype, subject_key, teacher_key, message_text, created_at) in enumerate(rows, start=1):
        name = escape(full_name or "Noma'lum")
        uname = f" @{escape(username)}" if username else ""
        line = f"{i}. <b>{name}</b>{uname}\n   ID: <code>{uid}</code>\n"
        if complaint_type == "teacher_complaint":
            line += f"   Kafedra: {escape(get_subject_name(subject_key or ''))}\n"
            line += f"   O'qituvchi: {escape(get_teacher_name(subject_key or '', teacher_key or ''))}\n"
        line += f"   Sana: {escape(created_at or '')}\n   Xabar: {escape(message_text or '')}"
        lines.append(line)
    text = "\n\n".join(lines)
    if len(text) > 4000:
        text = text[:4000] + "\n\n... qisqartirildi"
    return tr(user_id, text)


def export_complaints_filtered_to_excel(complaint_type: str) -> str:
    complaint_type = "suggestion" if complaint_type == "suggestion" else "teacher_complaint"
    filename = "suggestions_export.xlsx" if complaint_type == "suggestion" else "teacher_complaints_export.xlsx"
    path = os.path.join(DATA_DIR, filename)
    rows = execute_query_plain(
        """
        SELECT id, user_id, full_name, username, type, subject_key, teacher_key, message_text, created_at
        FROM complaints
        WHERE type = %s
        ORDER BY id DESC
        """,
        (complaint_type,), fetch='all'
    ) or []

    if Workbook is None:
        path = path.replace(".xlsx", ".csv")
        with open(path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(["ID", "Telegram ID", "F.I.Sh", "Username", "Kafedra", "O'qituvchi", "Sana", "Xabar"])
            for cid, uid, full_name, username, _ctype, subject_key, teacher_key, message_text, created_at in rows:
                writer.writerow([cid, uid, full_name or "", username or "", get_subject_name(subject_key or ""), get_teacher_name(subject_key or "", teacher_key or "") if teacher_key else "", created_at or "", message_text or ""])
        return path

    wb = Workbook()
    ws = wb.active
    ws.title = "Takliflar" if complaint_type == "suggestion" else "Shikoyatlar"
    ws_append_header(ws, ["ID", "Telegram ID", "F.I.Sh", "Username", "Kafedra", "O'qituvchi", "Sana", "Xabar"])
    for cid, uid, full_name, username, _ctype, subject_key, teacher_key, message_text, created_at in rows:
        ws.append([cid, uid, full_name or "", username or "", get_subject_name(subject_key or ""), get_teacher_name(subject_key or "", teacher_key or "") if teacher_key else "", created_at or "", message_text or ""])
    wb.save(path)
    return path

# =========================
# START / COMMANDS
# =========================
@dp.message(Command("start"))
async def start_handler(message: Message):
    user_id = message.from_user.id
    ensure_user(user_id)

    if has_access(user_id):
        await message.answer(get_home_text(user_id), parse_mode="HTML", reply_markup=home_keyboard(user_id))
        return

    if not await check_user_subscription(user_id):
        await message.answer(get_welcome_text(user_id), parse_mode="HTML", reply_markup=subscription_keyboard(user_id))
        return

    grant_access(user_id)
    await message.answer(get_home_text(user_id), parse_mode="HTML", reply_markup=home_keyboard(user_id))


@dp.message(Command("my_access"))
async def my_access_handler(message: Message):
    user_id = message.from_user.id
    ensure_user(user_id)
    if not is_admin(user_id):
        return
    row = execute_query_plain("SELECT access_granted FROM user_prefs WHERE user_id = %s", (user_id,), fetch='one')
    val = row[0] if row else "yo'q"
    await message.answer(f"access_granted: {val}")


@dp.message(Command("check_channel"))
async def check_channel_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    try:
        me = await bot.get_me()
        member = await bot.get_chat_member(CHANNEL_USERNAME, me.id)
        if member.status in {ChatMemberStatus.ADMINISTRATOR, ChatMemberStatus.CREATOR}:
            await message.answer(f"✅ Bot kanalni ko'ra olyapti va admin.\nKanal: {CHANNEL_USERNAME}")
        else:
            await message.answer("⚠️ Bot kanalni ko'ryapti, lekin admin emas. Obuna tekshiruvi to'liq ishlashi uchun botni kanalga admin qiling.")
    except Exception as e:
        await message.answer(f"❌ Kanalni tekshirib bo'lmadi.\nBotni kanalga admin qiling.\nXato: {e}")


@dp.message(Command("results"))
async def results_handler(message: Message):
    user_id = message.from_user.id
    ensure_user(user_id)
    await message.answer(get_results_menu_text(user_id, False), parse_mode="HTML", reply_markup=results_menu_keyboard_user(user_id))


@dp.message(Command("debug_eshnazarova"))
async def debug_eshnazarova_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        return

    rows = execute_query_plain(
        "SELECT user_id, full_name, username, subject_key, teacher_key, voted_at FROM votes WHERE teacher_key = 'aif_10' ORDER BY voted_at DESC",
        fetch='all'
    ) or []
    if not rows:
        await message.answer("Eshnazarova Maziya Allanazarovna uchun bazada ovoz yo'q.")
        return

    lines = ["Eshnazarova Maziya Allanazarovna uchun bazadagi ovozlar:"]
    for uid, full_name, username, subject_key, teacher_key, voted_at in rows:
        lines.append(f"ID: {uid} | {full_name or ''} | @{username or ''} | {subject_key}/{teacher_key} | {voted_at}")
    await message.answer("\n".join(lines[:50]))


@dp.message(Command("admin"))
async def admin_panel_handler(message: Message):
    user_id = message.from_user.id
    ensure_user(user_id)
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    await message.answer(get_admin_panel_text(user_id), parse_mode="HTML", reply_markup=admin_panel_keyboard(user_id))


@dp.message(Command("users"))
async def admin_users_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    text = await asyncio.to_thread(get_users_text, user_id)
    await message.answer(text, parse_mode="HTML", reply_markup=users_keyboard_admin(user_id))


@dp.message(Command("export"))
async def admin_export_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    filename = await asyncio.to_thread(export_votes_to_excel)
    await message.answer_document(FSInputFile(filename), caption="📁 Ovozlar Excel fayl ko'rinishida.")


@dp.message(Command("open"))
async def admin_open_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    open_voting()
    await message.answer("🟢 Ovoz berish ochildi.")


@dp.message(Command("close"))
async def admin_close_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    close_voting()
    await message.answer("🔴 Ovoz berish yopildi.")


@dp.message(Command("reset_votes"))
async def admin_reset_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    await message.answer(
        "⚠️ <b>Diqqat!</b>\n\nBarcha ovozlar o'chiriladi.\nDavom etasizmi?",
        parse_mode="HTML",
        reply_markup=reset_confirm_keyboard(user_id, "votes")
    )

# =========================
# USER CALLBACKS
# =========================
@dp.callback_query(F.data == "go_home")
async def go_home_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    ensure_user(user_id)
    WAITING_COMPLAINT_TEXT.discard(user_id)
    COMPLAINT_STATE.pop(user_id, None)
    if require_access_only(user_id):
        await safe_edit_message(callback, get_home_text(user_id), home_keyboard(user_id))
    else:
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "help_info")
async def help_info_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Orqaga"), callback_data="go_home"))
    await safe_edit_message(callback, get_help_text(user_id), kb.as_markup())
    await callback.answer()


@dp.callback_query(F.data == "my_vote")
async def my_vote_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not require_access_only(user_id):
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    await safe_edit_message(callback, get_my_vote_text(user_id), simple_back_home_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "my_ratings")
async def my_ratings_handler(callback: CallbackQuery):
    await callback.answer("Baholash funksiyasi olib tashlangan.", show_alert=True)


@dp.callback_query(F.data == "go_complaint_panel")
async def go_complaint_panel_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    ensure_user(user_id)

    if not has_access(user_id):
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return

    WAITING_COMPLAINT_TEXT.discard(user_id)
    COMPLAINT_STATE.pop(user_id, None)
    await safe_edit_message(callback, get_complaint_intro_text(user_id), complaint_menu_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "cancel_complaint")
async def cancel_complaint_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    WAITING_COMPLAINT_TEXT.discard(user_id)
    COMPLAINT_STATE.pop(user_id, None)

    if has_access(user_id):
        await safe_edit_message(callback, get_home_text(user_id), home_keyboard(user_id))
    else:
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
    await callback.answer(tr(user_id, "Bekor qilindi"))


@dp.callback_query(F.data == "complaint_teacher_start")
async def complaint_teacher_start_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not require_access_only(user_id):
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    WAITING_COMPLAINT_TEXT.discard(user_id)
    COMPLAINT_STATE[user_id] = {"mode": "teacher_complaint"}
    await safe_edit_message(callback, get_complaint_teacher_subject_text(user_id), complaint_subjects_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data.startswith("complaint_subject:"))
async def complaint_subject_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    subject_key = normalize_subject_key(callback.data.split(":", 1)[1])
    if subject_key not in get_subjects_from_db():
        await callback.answer(tr(user_id, "Noto'g'ri kafedra."), show_alert=True)
        return
    COMPLAINT_STATE[user_id] = {"mode": "teacher_complaint", "subject_key": subject_key}
    await safe_edit_message(callback, get_complaint_teacher_select_text(user_id, subject_key), complaint_teachers_keyboard(user_id, subject_key))
    await callback.answer()


@dp.callback_query(F.data.startswith("complaint_teacher:"))
async def complaint_teacher_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return
    _, subject_key, teacher_key = parts
    subject_key = normalize_subject_key(subject_key)
    if subject_key not in get_subjects_from_db() or teacher_key not in get_subjects_from_db().get(subject_key, {}).get("teachers", {}):
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return
    COMPLAINT_STATE[user_id] = {"mode": "teacher_complaint", "subject_key": subject_key, "teacher_key": teacher_key}
    WAITING_COMPLAINT_TEXT.add(user_id)
    await safe_edit_message(callback, get_complaint_write_text(user_id, "teacher_complaint", subject_key, teacher_key), complaint_cancel_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "suggestion_start")
async def suggestion_start_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not require_access_only(user_id):
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    COMPLAINT_STATE[user_id] = {"mode": "suggestion"}
    WAITING_COMPLAINT_TEXT.add(user_id)
    await safe_edit_message(callback, get_complaint_write_text(user_id, "suggestion"), complaint_cancel_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "confirm_complaint_send")
async def confirm_complaint_send_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    state = COMPLAINT_STATE.get(user_id)
    if not state or not state.get("text"):
        await callback.answer(tr(user_id, "Yuboriladigan matn topilmadi."), show_alert=True)
        return
    text = state.get("text", "")
    allowed, reason = await asyncio.to_thread(complaint_allowed, user_id, text, state.get("mode", "general"))
    if not allowed:
        await callback.answer(tr(user_id, reason), show_alert=True)
        return
    async with db_lock:
        await asyncio.to_thread(
            save_complaint,
            user_id=user_id,
            full_name=callback.from_user.full_name or "Noma'lum",
            username=callback.from_user.username or "",
            message_text=text,
            complaint_type=state.get("mode", "general"),
            subject_key=state.get("subject_key", ""),
            teacher_key=state.get("teacher_key", ""),
        )
    WAITING_COMPLAINT_TEXT.discard(user_id)
    COMPLAINT_STATE.pop(user_id, None)
    await safe_edit_message(callback, get_complaint_saved_text(user_id), home_keyboard(user_id))
    await callback.answer(tr(user_id, "Yuborildi"))


@dp.callback_query(F.data == "user_settings")
async def user_settings_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    ensure_user(user_id)
    await safe_edit_message(callback, get_settings_text(user_id), settings_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data.startswith("set_script:"))
async def set_script_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    ensure_user(user_id)

    script = callback.data.split(":", 1)[1]
    if script not in ("latin", "cyrillic"):
        await callback.answer("Xato", show_alert=True)
        return

    set_user_script(user_id, script)
    await safe_edit_message(callback, get_settings_text(user_id), settings_keyboard(user_id))
    await callback.answer(tr(user_id, "Yozuv turi saqlandi"))


@dp.callback_query(F.data == "check_subscription")
async def check_subscription_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    ensure_user(user_id)

    ok = await refresh_access_from_channel(user_id)
    if not ok:
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return

    await safe_edit_message(
        callback,
        "✅ <b>Obuna tasdiqlandi</b>\n\nEndi bosh menyudan bemalol foydalanishingiz mumkin:",
        home_keyboard(user_id)
    )
    await callback.answer(tr(user_id, "Tasdiqlandi"))


@dp.callback_query(F.data == "go_vote_panel")
async def go_vote_panel_handler(callback: CallbackQuery):
    user_id = callback.from_user.id

    if not await refresh_access_from_channel(user_id):
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    if has_voted(user_id):
        await safe_edit_message(callback, get_already_voted_text(user_id), home_keyboard(user_id))
        await callback.answer()
        return
    if not is_voting_open():
        await safe_edit_message(callback, get_closed_text(user_id), home_keyboard(user_id))
        await callback.answer()
        return
    await safe_edit_message(callback, get_subject_select_text(user_id), subjects_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data.startswith("subject:"))
async def subject_select_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not require_access_only(user_id):
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    if has_voted(user_id):
        await safe_edit_message(callback, get_already_voted_text(user_id), home_keyboard(user_id))
        await callback.answer()
        return
    if not is_voting_open():
        await safe_edit_message(callback, get_closed_text(user_id), home_keyboard(user_id))
        await callback.answer()
        return
    subject_key = normalize_subject_key(callback.data.split(":")[1])
    if subject_key not in get_subjects_from_db():
        await callback.answer("Noto'g'ri bo'lim tanlandi.", show_alert=True)
        return
    await safe_edit_message(callback, get_teacher_select_text(user_id, subject_key), teachers_keyboard(user_id, subject_key))
    await callback.answer()


@dp.callback_query(F.data.startswith("vote:"))
async def vote_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not require_access_only(user_id):
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    if not is_voting_open():
        await callback.answer(tr(user_id, "Hozir ovoz berish yopilgan."), show_alert=True)
        return
    if has_voted(user_id):
        await callback.answer(tr(user_id, "Siz faqat 1 marta ovoz bera olasiz."), show_alert=True)
        return
    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return
    _, subject_key, teacher_key = parts
    subject_key = normalize_subject_key(subject_key)
    if subject_key not in get_subjects_from_db() or teacher_key not in get_subjects_from_db().get(subject_key, {}).get("teachers", {}):
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return

    text = (
        f"❓ <b>Ovozingizni tasdiqlaysizmi?</b>\n\n"
        f"<b>Kafedra:</b> {get_subject_name(subject_key)}\n"
        f"<b>O'qituvchi:</b> {get_teacher_name(subject_key, teacher_key)}\n\n"
        f"Tasdiqlagandan keyin asosiy ovozni qayta o'zgartirib bo'lmaydi."
    )
    await safe_edit_message(callback, tr(user_id, text), confirm_vote_keyboard(user_id, subject_key, teacher_key))
    await callback.answer()


@dp.callback_query(F.data.startswith("confirm_vote:"))
async def confirm_vote_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not await refresh_access_from_channel(user_id):
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        return
    if not is_voting_open():
        await callback.answer(tr(user_id, "Hozir ovoz berish yopilgan."), show_alert=True)
        return
    if has_voted(user_id):
        await callback.answer(tr(user_id, "Siz faqat 1 marta ovoz bera olasiz."), show_alert=True)
        return

    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return
    _, subject_key, teacher_key = parts
    subject_key = normalize_subject_key(subject_key)
    if subject_key not in get_subjects_from_db() or teacher_key not in get_subjects_from_db().get(subject_key, {}).get("teachers", {}):
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except TelegramBadRequest:
        pass

    async with db_lock:
        saved = await asyncio.to_thread(
            save_vote,
            user_id=user_id,
            full_name=callback.from_user.full_name or "Noma'lum",
            username=callback.from_user.username or "",
            subject_key=subject_key,
            teacher_key=teacher_key
        )

    if not saved:
        await callback.answer(tr(user_id, "Siz allaqachon ovoz bergansiz."), show_alert=True)
        await safe_edit_message(callback, get_already_voted_text(user_id), home_keyboard(user_id))
        return

    text = (
        f"✅ <b>Ovoz muvaffaqiyatli qabul qilindi</b>\n\n"
        f"<b>Bo'lim:</b> {get_subject_name(subject_key)}\n"
        f"<b>Tanlovingiz:</b> {get_teacher_name(subject_key, teacher_key)}\n\n"
        f"Rahmat, sizning ovozingiz saqlandi."
    )
    await safe_edit_message(callback, tr(user_id, text), after_vote_keyboard(user_id))
    await callback.answer(tr(user_id, "Ovozingiz qabul qilindi!"))

# =========================
# RATING CALLBACKS
# =========================
@dp.callback_query(F.data == "go_rating_panel")
async def go_rating_panel_handler(callback: CallbackQuery):
    await callback.answer("Baholash funksiyasi olib tashlangan.", show_alert=True)

@dp.callback_query(F.data.startswith("rating_subject:"))
async def rating_subject_handler(callback: CallbackQuery):
    await callback.answer("Baholash funksiyasi olib tashlangan.", show_alert=True)

@dp.callback_query(F.data.startswith("rating_teacher:"))
async def rating_teacher_handler(callback: CallbackQuery):
    await callback.answer("Baholash funksiyasi olib tashlangan.", show_alert=True)

@dp.callback_query(F.data.startswith("rate:"))
async def rate_handler(callback: CallbackQuery):
    await callback.answer("Baholash funksiyasi olib tashlangan.", show_alert=True)

# =========================
# USER RESULTS
# =========================
@dp.callback_query(F.data == "show_results_menu_user")
async def show_results_menu_user(callback: CallbackQuery):
    user_id = callback.from_user.id
    await safe_edit_message(callback, get_results_menu_text(user_id, False), results_menu_keyboard_user(user_id))
    await callback.answer()


@dp.callback_query(F.data == "show_results_user:general")
async def show_results_user_general(callback: CallbackQuery):
    user_id = callback.from_user.id
    text = await asyncio.to_thread(get_general_results_text, user_id)
    text = add_refresh_time(text, user_id)
    await safe_edit_message(callback, text, results_keyboard_user(user_id, "general"))
    await callback.answer()


@dp.callback_query(F.data.startswith("show_results_user:"))
async def show_results_user(callback: CallbackQuery):
    user_id = callback.from_user.id
    scope = normalize_subject_key(callback.data.split(":", 1)[1].strip())

    if scope != "general" and scope not in get_subjects_from_db():
        await callback.answer(tr(user_id, "Noto'g'ri bo'lim."), show_alert=True)
        return

    text = await asyncio.to_thread(
        get_general_results_text if scope == "general" else get_subject_results_text,
        user_id,
        *(() if scope == "general" else (scope,))
    )
    text = add_refresh_time(text, user_id)

    await safe_edit_message(callback, text, results_keyboard_user(user_id, scope))
    await callback.answer()


@dp.callback_query(F.data == "refresh_results_user:general")
async def refresh_results_user_general(callback: CallbackQuery):
    user_id = callback.from_user.id
    refresh_key = "results_user:general"

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)
        text = await asyncio.to_thread(get_general_results_text, user_id)
        text = add_refresh_time(text, user_id)
        await safe_edit_message(callback, text, results_keyboard_user(user_id, "general"))
    finally:
        finish_refresh(user_id, refresh_key)


@dp.callback_query(F.data.startswith("refresh_results_user:"))
async def refresh_results_user(callback: CallbackQuery):
    user_id = callback.from_user.id
    scope = normalize_subject_key(callback.data.split(":", 1)[1].strip())
    refresh_key = f"results_user:{scope}"

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)

        if scope != "general" and scope not in get_subjects_from_db():
            await callback.answer(tr(user_id, "Noto'g'ri bo'lim."), show_alert=True)
            return

        text = await asyncio.to_thread(
            get_general_results_text if scope == "general" else get_subject_results_text,
            user_id,
            *(() if scope == "general" else (scope,))
        )
        text = add_refresh_time(text, user_id)

        await safe_edit_message(callback, text, results_keyboard_user(user_id, scope))
    finally:
        finish_refresh(user_id, refresh_key)


@dp.callback_query(F.data == "user_top_votes")
async def user_top_votes_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data="user_top_votes"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "📊 Natijalar"), callback_data="show_results_menu_user"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    text = await asyncio.to_thread(get_top_votes_text, user_id)
    text = add_refresh_time(text, user_id)
    await safe_edit_message(callback, text, kb.as_markup())
    await callback.answer()

# =========================
# ADMIN CALLBACKS
# =========================
@dp.callback_query(F.data == "back_admin_panel")
async def back_admin_panel_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, get_admin_panel_text(user_id), admin_panel_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "admin_cleanup_menu")
async def admin_cleanup_menu_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, "🧹 <b>Tozalash</b>\n\nQaysi ma'lumotlarni tozalamoqchisiz?", admin_cleanup_menu_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "admin_download_menu")
async def admin_download_menu_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, "📥 <b>Ma'lumotlarni yuklab olish</b>\n\nKerakli fayl turini tanlang:", admin_download_menu_keyboard(user_id))
    await callback.answer()


def get_admin_settings_text(user_id: int) -> str:
    status_icon = "🟢" if is_voting_open() else "🔴"
    status_text = "Ochiq" if is_voting_open() else "Yopiq"
    auto_icon = "🟢" if is_auto_voting_enabled() else "🔴"
    auto_text = "Yoqilgan" if is_auto_voting_enabled() else "O'chirilgan"
    return (
        "⚙️ <b>Sozlamalar</b>\n\n"
        f"{status_icon} <b>Ovoz berish holati:</b> {status_text}\n"
        f"{auto_icon} <b>Avtomatik rejim:</b> {auto_text}\n"
        f"⏰ <b>Avto vaqt:</b> {get_auto_voting_start()} - {get_auto_voting_end()}\n\n"
        "Kerakli bo'limni tanlang:"
    )


@dp.callback_query(F.data == "admin_settings_menu")
async def admin_settings_menu_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, get_admin_settings_text(user_id), admin_settings_menu_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "admin_open")
async def admin_open_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    open_voting()
    await safe_edit_message(callback, get_admin_settings_text(user_id), admin_settings_menu_keyboard(user_id))
    await callback.answer("Ovoz berish ochildi")


@dp.callback_query(F.data == "admin_close")
async def admin_close_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    close_voting()
    await safe_edit_message(callback, get_admin_settings_text(user_id), admin_settings_menu_keyboard(user_id))
    await callback.answer("Ovoz berish yopildi")


@dp.callback_query(F.data == "admin_auto_voting_toggle")
async def admin_auto_voting_toggle_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    enabled = not is_auto_voting_enabled()
    set_auto_voting_enabled(enabled)
    if enabled:
        apply_auto_voting_status()
    await safe_edit_message(callback, get_admin_settings_text(user_id), admin_settings_menu_keyboard(user_id))
    await callback.answer("Avtomatik rejim yoqildi" if enabled else "Avtomatik rejim o'chirildi")


@dp.callback_query(F.data == "admin_auto_voting_set_time")
async def admin_auto_voting_set_time_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    ADMIN_MANAGE_STATE[user_id] = {"action": "set_auto_voting_time"}
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text="⬅️ Sozlamalar", callback_data="admin_settings_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    await safe_edit_message(
        callback,
        "⏰ <b>Avtomatik ovoz vaqtini belgilang</b>\n\n"
        "Format: <code>09:00-18:00</code> yoki <code>09:00 18:00</code>\n"
        "Masalan: <code>08:30-17:45</code>\n\n"
        "Agar vaqt kechadan ertaga o'tsa, masalan <code>22:00-06:00</code>, bot shu oraliqda ochiq turadi.",
        kb.as_markup()
    )
    await callback.answer()


@dp.callback_query(F.data == "admin_results")
async def admin_results_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, get_results_menu_text(user_id, True), results_menu_keyboard_admin(user_id))
    await callback.answer()


@dp.callback_query(F.data == "show_results_admin:general")
async def show_results_admin_general(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    text = await asyncio.to_thread(get_general_results_text, user_id)
    text = add_refresh_time(text, user_id)
    await safe_edit_message(callback, text, results_keyboard_admin(user_id, "general"))
    await callback.answer()


@dp.callback_query(F.data == "refresh_results_admin:general")
async def refresh_results_admin_general(callback: CallbackQuery):
    user_id = callback.from_user.id
    refresh_key = "results_admin:general"

    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)
        text = await asyncio.to_thread(get_general_results_text, user_id)
        text = add_refresh_time(text, user_id)
        await safe_edit_message(callback, text, results_keyboard_admin(user_id, "general"))
    finally:
        finish_refresh(user_id, refresh_key)


@dp.callback_query(F.data.startswith("show_results_admin:"))
async def show_results_admin(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return

    scope = normalize_subject_key(callback.data.split(":", 1)[1].strip())
    if scope != "general" and scope not in get_subjects_from_db():
        await callback.answer(tr(user_id, "Noto'g'ri bo'lim."), show_alert=True)
        return

    text = await asyncio.to_thread(
        get_general_results_text if scope == "general" else get_subject_results_text,
        user_id,
        *(() if scope == "general" else (scope,))
    )
    text = add_refresh_time(text, user_id)

    await safe_edit_message(callback, text, results_keyboard_admin(user_id, scope))
    await callback.answer()


@dp.callback_query(F.data.startswith("refresh_results_admin:"))
async def refresh_results_admin_handler(callback: CallbackQuery):
    user_id = callback.from_user.id

    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return

    scope = normalize_subject_key(callback.data.split(":", 1)[1].strip())
    refresh_key = f"results_admin:{scope}"

    if scope != "general" and scope not in get_subjects_from_db():
        await callback.answer(tr(user_id, "Noto'g'ri bo'lim."), show_alert=True)
        return

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)
        text = await asyncio.to_thread(
            get_general_results_text if scope == "general" else get_subject_results_text,
            user_id,
            *(() if scope == "general" else (scope,))
        )
        text = add_refresh_time(text, user_id)

        await safe_edit_message(callback, text, results_keyboard_admin(user_id, scope))
    finally:
        finish_refresh(user_id, refresh_key)


@dp.callback_query(F.data == "admin_rating_stats")
async def admin_rating_stats_callback(callback: CallbackQuery):
    await callback.answer("Baholash funksiyasi olib tashlangan.", show_alert=True)

@dp.callback_query(F.data.startswith("show_rating_stats:"))
async def show_rating_stats_callback(callback: CallbackQuery):
    await callback.answer("Baholash funksiyasi olib tashlangan.", show_alert=True)

@dp.callback_query(F.data.startswith("refresh_rating_stats:"))
async def refresh_rating_stats_callback(callback: CallbackQuery):
    await callback.answer("Baholash funksiyasi olib tashlangan.", show_alert=True)

@dp.callback_query(F.data == "admin_top_ratings")
async def admin_top_ratings_callback(callback: CallbackQuery):
    await callback.answer("Baholash funksiyasi olib tashlangan.", show_alert=True)

@dp.callback_query(F.data == "admin_top_votes")
async def admin_top_votes_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data="admin_top_votes"))
    kb.row(InlineKeyboardButton(text="⬅️ TOP menyusi", callback_data="admin_top_votes_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    text = await asyncio.to_thread(get_top_votes_text, user_id)
    text = add_refresh_time(text, user_id)
    await safe_edit_message(callback, text, kb.as_markup())
    await callback.answer()


@dp.callback_query(F.data == "admin_teacher_stats")
async def admin_teacher_stats_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, tr(user_id, "👤 <b>O'qituvchi statistikasi</b>\n\nKafedrani tanlang:"), teacher_stats_subjects_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data.startswith("teacher_stats_subject:"))
async def teacher_stats_subject_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    subject_key = normalize_subject_key(callback.data.split(":", 1)[1])
    if subject_key not in get_subjects_from_db():
        await callback.answer(tr(user_id, "Noto'g'ri kafedra."), show_alert=True)
        return
    await safe_edit_message(
        callback,
        tr(user_id, f"👤 <b>{get_subject_name(subject_key)}</b>\n\nStatistika uchun o'qituvchini tanlang:"),
        teacher_stats_teachers_keyboard(user_id, subject_key)
    )
    await callback.answer()


@dp.callback_query(F.data.startswith("teacher_stats:"))
async def teacher_stats_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return
    _, subject_key, teacher_key = parts
    subject_key = normalize_subject_key(subject_key)
    text = await asyncio.to_thread(get_teacher_detailed_stats_text, user_id, subject_key, teacher_key)
    await safe_edit_message(callback, text, teacher_stats_keyboard(user_id, subject_key))
    await callback.answer()


@dp.callback_query(F.data == "admin_backup")
async def admin_backup_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await callback.answer("Backup tayyorlanmoqda...")
    try:
        backup_path = await asyncio.to_thread(create_backup_zip)
        await callback.message.answer_document(
            FSInputFile(backup_path),
            caption="💾 Backup: ovozlar Excel va shikoyatlar fayli."
        )
    except Exception as e:
        logging.error(f"Backup yaratishda xatolik: {e}")
        await callback.message.answer(f"❌ Backup yaratishda xatolik: {e}")


@dp.callback_query(F.data == "admin_complaints_menu")
async def admin_complaints_menu_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, "📩 <b>Shikoyat va takliflar</b>\n\nKerakli bo'limni tanlang:", admin_complaints_menu_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data.in_({"admin_complaints_teacher", "admin_complaints_suggestions"}))
async def admin_complaints_filtered_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    ctype = "suggestion" if callback.data == "admin_complaints_suggestions" else "teacher_complaint"
    text = await asyncio.to_thread(get_complaints_text_filtered, user_id, ctype)
    text = add_refresh_time(text, user_id)
    await safe_edit_message(callback, text, complaints_keyboard_admin(user_id, ctype))
    await callback.answer()


@dp.callback_query(F.data.startswith("refresh_admin_complaints:"))
async def refresh_admin_complaints_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    ctype = callback.data.split(":", 1)[1]
    refresh_key = f"admin_complaints:{ctype}"

    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)
        text = await asyncio.to_thread(get_complaints_text_filtered, user_id, ctype)
        text = add_refresh_time(text, user_id)
        await safe_edit_message(callback, text, complaints_keyboard_admin(user_id, ctype))
    finally:
        finish_refresh(user_id, refresh_key)


@dp.callback_query(F.data == "admin_export_complaints_docx")
async def admin_export_complaints_docx_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    filename = await asyncio.to_thread(export_complaints_to_docx)
    await callback.message.answer_document(
        FSInputFile(filename),
        caption="📄 Shikoyat va takliflar Word fayl ko'rinishida."
    )
    await callback.answer()


@dp.callback_query(F.data == "admin_export_teacher_complaints_excel")
async def admin_export_teacher_complaints_excel_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    filename = await asyncio.to_thread(export_complaints_filtered_to_excel, "teacher_complaint")
    await callback.message.answer_document(FSInputFile(filename), caption="👤 O'qituvchi ustidan shikoyatlar fayli.")
    await callback.answer()


@dp.callback_query(F.data == "admin_export_suggestions_excel")
async def admin_export_suggestions_excel_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    filename = await asyncio.to_thread(export_complaints_filtered_to_excel, "suggestion")
    await callback.message.answer_document(FSInputFile(filename), caption="💡 Takliflar fayli.")
    await callback.answer()


@dp.callback_query(F.data == "admin_export_subjects_excel")
async def admin_export_subjects_excel_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    filename = await asyncio.to_thread(export_subjects_ranking_to_excel)
    await callback.message.answer_document(FSInputFile(filename), caption="🏫 Kafedralar natijasi Excel fayli.")
    await callback.answer()


@dp.callback_query(F.data == "admin_users")
async def admin_users_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    text = await asyncio.to_thread(get_users_text, user_id)
    text = add_refresh_time(text, user_id)
    await safe_edit_message(callback, text, users_keyboard_admin(user_id))
    await callback.answer()


@dp.callback_query(F.data == "refresh_admin_users")
async def refresh_admin_users(callback: CallbackQuery):
    user_id = callback.from_user.id
    refresh_key = "admin_users"

    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)
        text = await asyncio.to_thread(get_users_text, user_id)
        text = add_refresh_time(text, user_id)
        await safe_edit_message(callback, text, users_keyboard_admin(user_id))
    finally:
        finish_refresh(user_id, refresh_key)


@dp.callback_query(F.data == "admin_export_votes_excel")
async def admin_export_votes_excel_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    filename = await asyncio.to_thread(export_votes_to_excel)
    await callback.message.answer_document(FSInputFile(filename), caption="📁 Ovozlar Excel fayl ko'rinishida.")
    await callback.answer()


@dp.callback_query(F.data == "admin_export_users_excel")
async def admin_export_users_excel_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    filename = await asyncio.to_thread(export_users_to_excel)
    await callback.message.answer_document(FSInputFile(filename), caption="👥 Foydalanuvchilar Excel fayli.")
    await callback.answer()


@dp.callback_query(F.data == "admin_export_rating_excel")
async def admin_export_rating_excel_callback(callback: CallbackQuery):
    await callback.answer("Baholash funksiyasi olib tashlangan.", show_alert=True)


@dp.callback_query(F.data == "admin_reset_rating_confirm")
async def admin_reset_rating_confirm_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, "⚠️ <b>Diqqat!</b>\n\nBarcha rating baholari o'chiriladi.\nDavom etasizmi?", reset_confirm_keyboard(user_id, "rating"))
    await callback.answer()


@dp.callback_query(F.data == "admin_reset_votes_confirm")
async def admin_reset_votes_confirm_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(
        callback,
        "⚠️ <b>Diqqat!</b>\n\nBarcha ovozlar o'chiriladi.\nDavom etasizmi?",
        reset_confirm_keyboard(user_id, "votes")
    )
    await callback.answer()


@dp.callback_query(F.data == "admin_reset_complaints_confirm")
async def admin_reset_complaints_confirm_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(
        callback,
        "⚠️ <b>Diqqat!</b>\n\nBarcha shikoyat va takliflar o'chiriladi.\nDavom etasizmi?",
        reset_confirm_keyboard(user_id, "complaints")
    )
    await callback.answer()


@dp.callback_query(F.data == "admin_reset_complaints")
async def admin_reset_complaints_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    async with db_lock:
        await asyncio.to_thread(reset_complaints)
    await safe_edit_message(callback, get_admin_panel_text(user_id), admin_panel_keyboard(user_id))
    await callback.answer("Shikoyat va takliflar tozalandi!")


@dp.callback_query(F.data == "cancel_reset")
async def cancel_reset_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, get_admin_panel_text(user_id), admin_panel_keyboard(user_id))
    await callback.answer(tr(user_id, "Bekor qilindi"))


@dp.callback_query(F.data == "admin_reset_votes")
async def admin_reset_votes_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    async with db_lock:
        await asyncio.to_thread(reset_votes)
    await safe_edit_message(callback, get_admin_panel_text(user_id), admin_panel_keyboard(user_id))
    await callback.answer("Ovozlar reset qilindi!")


@dp.callback_query(F.data == "admin_reset_rating")
async def admin_reset_rating_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    async with db_lock:
        await asyncio.to_thread(reset_ratings)
    await safe_edit_message(callback, get_admin_panel_text(user_id), admin_panel_keyboard(user_id))
    await callback.answer("Rating reset qilindi!")


@dp.callback_query(F.data == "admin_export")
async def admin_export_callback(callback: CallbackQuery):
    await admin_export_votes_excel_callback(callback)


@dp.callback_query(F.data == "admin_reset_confirm")
async def admin_reset_confirm_old_callback(callback: CallbackQuery):
    await admin_reset_votes_confirm_callback(callback)


@dp.callback_query(F.data == "admin_reset")
async def admin_reset_old_callback(callback: CallbackQuery):
    await admin_reset_votes_callback(callback)

# =========================
# ADMIN DB MANAGE FUNCTIONS
# =========================
def db_add_subject(subject_key: str, subject_name: str) -> bool:
    try:
        row = execute_query_plain("SELECT MAX(sort_order) FROM db_subjects", fetch='one')
        max_order = (row[0] or 0) + 1
        execute_query(
            "INSERT INTO db_subjects (subject_key, subject_name, sort_order) VALUES ($1, $2, $3)",
            (subject_key, subject_name, max_order)
        )
        invalidate_subjects_cache()
        return True
    except PgIntegrityError:
        return False


def db_edit_subject(subject_key: str, new_name: str) -> bool:
    rowcount = execute_query("UPDATE db_subjects SET subject_name = $1 WHERE subject_key = $2", (new_name, subject_key))
    invalidate_subjects_cache()
    return rowcount > 0


def db_delete_subject(subject_key: str) -> bool:
    execute_query("DELETE FROM db_teachers WHERE subject_key = $1", (subject_key,))
    execute_query("DELETE FROM db_subjects WHERE subject_key = $1", (subject_key,))
    invalidate_subjects_cache()
    return True


def db_add_teacher(subject_key: str, teacher_key: str, teacher_name: str) -> bool:
    try:
        execute_query(
            "INSERT INTO db_teachers (teacher_key, subject_key, teacher_name) VALUES ($1, $2, $3)",
            (teacher_key, subject_key, teacher_name)
        )
        invalidate_subjects_cache()
        return True
    except PgIntegrityError:
        return False


def db_edit_teacher(subject_key: str, teacher_key: str, new_name: str) -> bool:
    rowcount = execute_query(
        "UPDATE db_teachers SET teacher_name = $1 WHERE subject_key = $2 AND teacher_key = $3",
        (new_name, subject_key, teacher_key)
    )
    invalidate_subjects_cache()
    return rowcount > 0


def db_delete_teacher(subject_key: str, teacher_key: str) -> bool:
    rowcount = execute_query("DELETE FROM db_teachers WHERE subject_key = $1 AND teacher_key = $2", (subject_key, teacher_key))
    invalidate_subjects_cache()
    return rowcount > 0


def generate_teacher_key(subject_key: str) -> str:
    prefix = subject_key[:3]
    subjects = get_subjects_from_db()
    existing = set(subjects.get(subject_key, {}).get("teachers", {}).keys())
    i = 1
    while f"{prefix}_{i}" in existing:
        i += 1
    return f"{prefix}_{i}"


def generate_subject_key(name: str) -> str:
    import re as _re
    base = _re.sub(r"[^a-zA-Z0-9]", "", name.lower().replace(" ", "_"))[:8] or "subj"
    existing = set(get_subjects_from_db().keys())
    key = base
    i = 1
    while key in existing:
        key = f"{base}{i}"
        i += 1
    return key

# =========================
# ADMIN MANAGE KEYBOARDS
# =========================
def admin_manage_menu_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text="🏫 Bo'limlarni boshqarish", callback_data="manage_subjects"),
        InlineKeyboardButton(text="👨‍🏫 O'qituvchilarni boshqarish", callback_data="manage_teachers_select_subject")
    )
    kb.row(InlineKeyboardButton(text="⬅️ Sozlamalar", callback_data="admin_settings_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def manage_subjects_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    subjects = get_subjects_from_db()
    for skey, sdata in subjects.items():
        kb.row(InlineKeyboardButton(text=f"📂 {sdata['name']}", callback_data=f"manage_subject_actions:{skey}"))
    kb.row(InlineKeyboardButton(text="➕ Yangi bo'lim qo'shish", callback_data="manage_subject_add"))
    kb.row(InlineKeyboardButton(text="⬅️ Boshqarish menyusi", callback_data="admin_manage_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Sozlamalar", callback_data="admin_settings_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def manage_subject_actions_keyboard(user_id: int, subject_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text="✏️ Nomini o'zgartirish", callback_data=f"manage_subject_edit:{subject_key}"),
        InlineKeyboardButton(text="🗑 O'chirish", callback_data=f"manage_subject_delete_confirm:{subject_key}")
    )
    kb.row(InlineKeyboardButton(text="⬅️ Bo'limlar", callback_data="manage_subjects"))
    kb.row(InlineKeyboardButton(text="⬅️ Boshqarish menyusi", callback_data="admin_manage_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Sozlamalar", callback_data="admin_settings_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def manage_teachers_select_subject_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    subjects = get_subjects_from_db()
    for skey, sdata in subjects.items():
        kb.row(InlineKeyboardButton(text=sdata['name'], callback_data=f"manage_teachers:{skey}"))
    kb.row(InlineKeyboardButton(text="⬅️ Boshqarish menyusi", callback_data="admin_manage_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Sozlamalar", callback_data="admin_settings_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def manage_teachers_keyboard(user_id: int, subject_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    subjects = get_subjects_from_db()
    teachers = subjects.get(subject_key, {}).get("teachers", {})
    for tkey, tname in teachers.items():
        kb.row(InlineKeyboardButton(text=tname, callback_data=f"manage_teacher_actions:{subject_key}:{tkey}"))
    kb.row(InlineKeyboardButton(text="➕ Yangi o'qituvchi qo'shish", callback_data=f"manage_teacher_add:{subject_key}"))
    kb.row(InlineKeyboardButton(text="⬅️ Bo'limlar", callback_data="manage_teachers_select_subject"))
    kb.row(InlineKeyboardButton(text="⬅️ Boshqarish menyusi", callback_data="admin_manage_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Sozlamalar", callback_data="admin_settings_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def manage_teacher_actions_keyboard(user_id: int, subject_key: str, teacher_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text="✏️ Nomini o'zgartirish", callback_data=f"manage_teacher_edit:{subject_key}:{teacher_key}"),
        InlineKeyboardButton(text="👥 O'quvchilar soni", callback_data=f"manage_teacher_students:{subject_key}:{teacher_key}")
    )
    kb.row(InlineKeyboardButton(text="🗑 O'chirish", callback_data=f"manage_teacher_delete_confirm:{subject_key}:{teacher_key}"))
    kb.row(InlineKeyboardButton(text="⬅️ O'qituvchilar", callback_data=f"manage_teachers:{subject_key}"))
    kb.row(InlineKeyboardButton(text="⬅️ Bo'limlar", callback_data="manage_teachers_select_subject"))
    kb.row(InlineKeyboardButton(text="⬅️ Boshqarish menyusi", callback_data="admin_manage_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Sozlamalar", callback_data="admin_settings_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def manage_cancel_keyboard(user_id: int, back_cb: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text="⬅️ Orqaga / Bekor qilish", callback_data=back_cb))
    kb.row(InlineKeyboardButton(text="⬅️ Sozlamalar", callback_data="admin_settings_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def top_votes_menu_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text="👨‍🏫 O'qituvchilar TOP 10", callback_data="admin_top_votes"),
        InlineKeyboardButton(text="🏫 Bo'limlar TOP 10", callback_data="admin_top_votes_by_subject")
    )
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()

# =========================
# ADMIN TOP VOTES MENU CALLBACKS
# =========================
@dp.callback_query(F.data == "admin_top_votes_menu")
async def admin_top_votes_menu_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, "🥇 <b>TOP 10 ovoz bo'yicha</b>\n\nQaysi ko'rinishda ko'rmoqchisiz?", top_votes_menu_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "admin_top_votes_by_subject")
async def admin_top_votes_by_subject_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data="admin_top_votes_by_subject"))
    kb.row(InlineKeyboardButton(text="⬅️ TOP menyusi", callback_data="admin_top_votes_menu"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    text = await asyncio.to_thread(get_top_votes_by_subject_text, user_id)
    text = add_refresh_time(text, user_id)
    await safe_edit_message(callback, text, kb.as_markup())
    await callback.answer()

# =========================
# ADMIN MANAGE CALLBACKS
# =========================
@dp.callback_query(F.data == "admin_manage_menu")
async def admin_manage_menu_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, "⚙️ <b>Boshqarish menyusi</b>\n\nNimani boshqarmoqchisiz?", admin_manage_menu_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "manage_subjects")
async def manage_subjects_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    subjects = get_subjects_from_db()
    text = f"🏫 <b>Bo'limlar ro'yxati</b>\n\nJami: {len(subjects)} ta bo'lim"
    await safe_edit_message(callback, text, manage_subjects_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data.startswith("manage_subject_actions:"))
async def manage_subject_actions_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    subject_key = callback.data.split(":", 1)[1]
    sname = get_subject_name(subject_key)
    subjects = get_subjects_from_db()
    teacher_count = len(subjects.get(subject_key, {}).get("teachers", {}))
    text = f"📂 <b>{sname}</b>\n\nO'qituvchilar soni: {teacher_count}\n\nNimani qilmoqchisiz?"
    await safe_edit_message(callback, text, manage_subject_actions_keyboard(user_id, subject_key))
    await callback.answer()


@dp.callback_query(F.data == "manage_subject_add")
async def manage_subject_add_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    ADMIN_MANAGE_STATE[user_id] = {"action": "add_subject"}
    await safe_edit_message(callback,
        "➕ <b>Yangi bo'lim qo'shish</b>\n\nBo'lim nomini yozing (masalan: Tarix fanlari):",
        manage_cancel_keyboard(user_id, "manage_subjects"))
    await callback.answer()


@dp.callback_query(F.data.startswith("manage_subject_edit:"))
async def manage_subject_edit_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    subject_key = callback.data.split(":", 1)[1]
    sname = get_subject_name(subject_key)
    ADMIN_MANAGE_STATE[user_id] = {"action": "edit_subject", "subject_key": subject_key}
    await safe_edit_message(callback,
        f"✏️ <b>Bo'lim nomini o'zgartirish</b>\n\nHozirgi nom: <b>{sname}</b>\n\nYangi nomni yozing:",
        manage_cancel_keyboard(user_id, f"manage_subject_actions:{subject_key}"))
    await callback.answer()


@dp.callback_query(F.data.startswith("manage_subject_delete_confirm:"))
async def manage_subject_delete_confirm_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    subject_key = callback.data.split(":", 1)[1]
    sname = get_subject_name(subject_key)
    subjects = get_subjects_from_db()
    teacher_count = len(subjects.get(subject_key, {}).get("teachers", {}))
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text="❌ Bekor qilish", callback_data=f"manage_subject_actions:{subject_key}"),
        InlineKeyboardButton(text="✅ Ha, o'chirish", callback_data=f"manage_subject_delete:{subject_key}")
    )
    kb.row(InlineKeyboardButton(text="⬅️ Bo'limlar", callback_data="manage_subjects"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    await safe_edit_message(callback,
        f"⚠️ <b>Diqqat!</b>\n\n<b>{sname}</b> bo'limi va unga tegishli <b>{teacher_count} ta o'qituvchi</b> o'chiriladi!\n\nDavom etasizmi?",
        kb.as_markup())
    await callback.answer()


@dp.callback_query(F.data.startswith("manage_subject_delete:"))
async def manage_subject_delete_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    subject_key = callback.data.split(":", 1)[1]
    sname = get_subject_name(subject_key)
    async with db_lock:
        await asyncio.to_thread(db_delete_subject, subject_key)
    subjects = get_subjects_from_db()
    text = f"✅ <b>{sname}</b> bo'limi o'chirildi.\n\nJami: {len(subjects)} ta bo'lim"
    await safe_edit_message(callback, text, manage_subjects_keyboard(user_id))
    await callback.answer("O'chirildi!")


@dp.callback_query(F.data == "manage_teachers_select_subject")
async def manage_teachers_select_subject_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, "👨‍🏫 <b>O'qituvchilarni boshqarish</b>\n\nBo'limni tanlang:", manage_teachers_select_subject_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data.startswith("manage_teachers:"))
async def manage_teachers_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    subject_key = callback.data.split(":", 1)[1]
    sname = get_subject_name(subject_key)
    subjects = get_subjects_from_db()
    teacher_count = len(subjects.get(subject_key, {}).get("teachers", {}))
    text = f"👨‍🏫 <b>{sname}</b>\n\nO'qituvchilar soni: {teacher_count}\n\nO'qituvchini tanlang:"
    await safe_edit_message(callback, text, manage_teachers_keyboard(user_id, subject_key))
    await callback.answer()


@dp.callback_query(F.data.startswith("manage_teacher_actions:"))
async def manage_teacher_actions_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer("Xato", show_alert=True)
        return
    _, subject_key, teacher_key = parts
    tname = get_teacher_name(subject_key, teacher_key)
    sname = get_subject_name(subject_key)
    text = f"👤 <b>{tname}</b>\n📂 {sname}\n\nNimani qilmoqchisiz?"
    await safe_edit_message(callback, text, manage_teacher_actions_keyboard(user_id, subject_key, teacher_key))
    await callback.answer()


@dp.callback_query(F.data.startswith("manage_teacher_add:"))
async def manage_teacher_add_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    subject_key = callback.data.split(":", 1)[1]
    sname = get_subject_name(subject_key)
    ADMIN_MANAGE_STATE[user_id] = {"action": "add_teacher", "subject_key": subject_key}
    await safe_edit_message(callback,
        f"➕ <b>Yangi o'qituvchi qo'shish</b>\n📂 Bo'lim: <b>{sname}</b>\n\nO'qituvchi to'liq ismini yozing (F.I.Sh):",
        manage_cancel_keyboard(user_id, f"manage_teachers:{subject_key}"))
    await callback.answer()


@dp.callback_query(F.data.startswith("manage_teacher_edit:"))
async def manage_teacher_edit_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer("Xato", show_alert=True)
        return
    _, subject_key, teacher_key = parts
    tname = get_teacher_name(subject_key, teacher_key)
    ADMIN_MANAGE_STATE[user_id] = {"action": "edit_teacher", "subject_key": subject_key, "teacher_key": teacher_key}
    await safe_edit_message(callback,
        f"✏️ <b>O'qituvchi nomini o'zgartirish</b>\n\nHozirgi ism: <b>{tname}</b>\n\nYangi to'liq ismni yozing:",
        manage_cancel_keyboard(user_id, f"manage_teacher_actions:{subject_key}:{teacher_key}"))
    await callback.answer()


@dp.callback_query(F.data.startswith("manage_teacher_students:"))
async def manage_teacher_students_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer("Xato", show_alert=True)
        return
    _, subject_key, teacher_key = parts
    tname = get_teacher_name(subject_key, teacher_key)
    current = get_teacher_student_count(subject_key, teacher_key)
    ADMIN_MANAGE_STATE[user_id] = {"action": "set_teacher_students", "subject_key": subject_key, "teacher_key": teacher_key}
    await safe_edit_message(
        callback,
        f"👥 <b>O'quvchilar sonini kiritish</b>\n\nO'qituvchi: <b>{tname}</b>\nHozirgi son: <b>{current}</b>\n\nYangi sonni raqam bilan yuboring:",
        manage_cancel_keyboard(user_id, f"manage_teacher_actions:{subject_key}:{teacher_key}")
    )
    await callback.answer()


@dp.callback_query(F.data.startswith("manage_teacher_delete_confirm:"))
async def manage_teacher_delete_confirm_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer("Xato", show_alert=True)
        return
    _, subject_key, teacher_key = parts
    tname = get_teacher_name(subject_key, teacher_key)
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text="❌ Bekor qilish", callback_data=f"manage_teacher_actions:{subject_key}:{teacher_key}"),
        InlineKeyboardButton(text="✅ Ha, o'chirish", callback_data=f"manage_teacher_delete:{subject_key}:{teacher_key}")
    )
    kb.row(InlineKeyboardButton(text="⬅️ O'qituvchilar", callback_data=f"manage_teachers:{subject_key}"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    await safe_edit_message(callback,
        f"⚠️ <b>Diqqat!</b>\n\n<b>{tname}</b> o'chiriladi!\n\nDavom etasizmi?",
        kb.as_markup())
    await callback.answer()


@dp.callback_query(F.data.startswith("manage_teacher_delete:"))
async def manage_teacher_delete_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer("Xato", show_alert=True)
        return
    _, subject_key, teacher_key = parts
    tname = get_teacher_name(subject_key, teacher_key)
    async with db_lock:
        await asyncio.to_thread(db_delete_teacher, subject_key, teacher_key)
    await safe_edit_message(callback,
        f"✅ <b>{tname}</b> o'chirildi.",
        manage_teachers_keyboard(user_id, subject_key))
    await callback.answer("O'chirildi!")

# =========================
# TEXT HANDLER
# =========================
@dp.message(F.text)
async def text_handler(message: Message):
    user_id = message.from_user.id
    ensure_user(user_id)

    if is_admin(user_id) and user_id in ADMIN_MANAGE_STATE:
        state = ADMIN_MANAGE_STATE.pop(user_id)
        text = (message.text or "").strip()
        action = state.get("action")

        if action == "add_subject":
            if not text:
                await message.answer("Bo'sh nom kiritilmadi.")
                return
            skey = generate_subject_key(text)
            async with db_lock:
                ok = await asyncio.to_thread(db_add_subject, skey, text)
            if ok:
                await message.answer(
                    f"✅ <b>{text}</b> bo'limi qo'shildi!", parse_mode="HTML",
                    reply_markup=manage_subjects_keyboard(user_id)
                )
            else:
                await message.answer("❌ Xatolik: bunday kalit allaqachon mavjud.",
                    reply_markup=manage_subjects_keyboard(user_id))
            return

        elif action == "edit_subject":
            subject_key = state.get("subject_key")
            if not text:
                await message.answer("Bo'sh nom kiritilmadi.")
                return
            async with db_lock:
                await asyncio.to_thread(db_edit_subject, subject_key, text)
            await message.answer(
                f"✅ Bo'lim nomi <b>{text}</b> ga o'zgartirildi!", parse_mode="HTML",
                reply_markup=manage_subjects_keyboard(user_id)
            )
            return

        elif action == "add_teacher":
            subject_key = state.get("subject_key")
            if not text:
                await message.answer("Bo'sh ism kiritilmadi.")
                return
            tkey = generate_teacher_key(subject_key)
            async with db_lock:
                await asyncio.to_thread(db_add_teacher, subject_key, tkey, text)
            await message.answer(
                f"✅ <b>{text}</b> o'qituvchi sifatida qo'shildi!", parse_mode="HTML",
                reply_markup=manage_teachers_keyboard(user_id, subject_key)
            )
            return

        elif action == "edit_teacher":
            subject_key = state.get("subject_key")
            teacher_key = state.get("teacher_key")
            if not text:
                await message.answer("Bo'sh ism kiritilmadi.")
                return
            async with db_lock:
                await asyncio.to_thread(db_edit_teacher, subject_key, teacher_key, text)
            await message.answer(
                f"✅ O'qituvchi ismi <b>{text}</b> ga o'zgartirildi!", parse_mode="HTML",
                reply_markup=manage_teachers_keyboard(user_id, subject_key)
            )
            return

        elif action == "set_teacher_students":
            subject_key = state.get("subject_key")
            teacher_key = state.get("teacher_key")
            if not text.isdigit():
                ADMIN_MANAGE_STATE[user_id] = state
                await message.answer("❌ Faqat raqam kiriting. Masalan: 120", reply_markup=manage_cancel_keyboard(user_id, f"manage_teacher_actions:{subject_key}:{teacher_key}"))
                return
            student_count = int(text)
            async with db_lock:
                await asyncio.to_thread(set_teacher_student_count, subject_key, teacher_key, student_count)
            await message.answer(
                f"✅ O'quvchilar soni <b>{student_count}</b> qilib saqlandi!", parse_mode="HTML",
                reply_markup=manage_teacher_actions_keyboard(user_id, subject_key, teacher_key)
            )
            return

        elif action == "set_auto_voting_time":
            cleaned = text.replace("—", "-").replace("–", "-").replace("dan", " ").replace("gacha", " ")
            parts = re.findall(r"([01]?\d|2[0-3]):([0-5]\d)", cleaned)
            if len(parts) < 2:
                ADMIN_MANAGE_STATE[user_id] = state
                await message.answer(
                    "❌ Format noto'g'ri. Masalan: <code>09:00-18:00</code>",
                    parse_mode="HTML",
                    reply_markup=manage_cancel_keyboard(user_id, "admin_settings_menu")
                )
                return
            start_time = f"{int(parts[0][0]):02d}:{parts[0][1]}"
            end_time = f"{int(parts[1][0]):02d}:{parts[1][1]}"
            async with db_lock:
                await asyncio.to_thread(set_auto_voting_schedule, start_time, end_time)
                if await asyncio.to_thread(is_auto_voting_enabled):
                    await asyncio.to_thread(apply_auto_voting_status)
            await message.answer(
                f"✅ Avtomatik vaqt saqlandi: <b>{start_time} - {end_time}</b>",
                parse_mode="HTML",
                reply_markup=admin_settings_menu_keyboard(user_id)
            )
            return

    if user_id in WAITING_COMPLAINT_TEXT:
        text = (message.text or "").strip()
        if not text:
            await message.answer("Bo'sh xabar qabul qilinmaydi. Iltimos, matn yozing.")
            return

        if len(text) > COMPLAINT_MAX_LENGTH:
            await message.answer(tr(user_id, f"Xabar juda uzun. Iltimos, {COMPLAINT_MAX_LENGTH} ta belgidan oshirmang."), parse_mode="HTML", reply_markup=complaint_cancel_keyboard(user_id))
            return

        WAITING_COMPLAINT_TEXT.discard(user_id)
        state = COMPLAINT_STATE.get(user_id, {"mode": "suggestion"})
        state["text"] = text
        COMPLAINT_STATE[user_id] = state
        await message.answer(
            get_complaint_confirm_text(user_id, state),
            parse_mode="HTML",
            reply_markup=complaint_confirm_keyboard(user_id)
        )
        return

    if message.text and message.text.lower() == "results":
        await message.answer(
            get_results_menu_text(user_id, False),
            parse_mode="HTML",
            reply_markup=results_menu_keyboard_user(user_id)
        )

# =========================
# MAIN
# =========================
async def on_startup():
    await init_pool(DATABASE_URL, min_size=5, max_size=20, command_timeout=30)
    await asyncio.to_thread(init_db)
    asyncio.create_task(auto_voting_scheduler())
    logging.info("Bot ishga tushdi. PostgreSQL asyncpg pool, cache va indexes tayyor.")


async def main():
    await on_startup()
    try:
        await dp.start_polling(bot)
    finally:
        await close_pool()

if __name__ == "__main__":
    asyncio.run(main())
